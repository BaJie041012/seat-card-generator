# ==============================================================================
# 文件名称: gui.py
# 功能描述: GUI用户界面模块，基于tkinter提供图形化操作界面
# 创建日期: 2026-03-14
# 作    者: 戒者有八
# ==============================================================================
"""
GUI用户界面 - 基于tkinter的图形界面

本模块提供完整的图形用户界面，包括:
    - 文本输入和编辑区域
    - AI智能提取功能
    - 模板选择和管理
    - 结果展示和导出
    - 系统设置对话框

设计思路:
    使用tkinter标准库构建跨平台GUI界面。
    采用面向对象设计，将界面组件和业务逻辑封装在类中。
    使用多线程处理AI请求，避免界面卡顿。
"""

import os
import threading
from typing import Optional, List
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

from config import CONFIG, AIConfig
from ai_service import AIService, create_ai_service
from template_processor import TemplateProcessor, create_template_processor
from text_extractor import AIExtractor, create_extractor



# ------------------------------------------------------------------------------
# 主应用程序类
# 功能: 应用程序主窗口，整合所有功能模块
# 设计思路: MVC模式，分离界面、业务逻辑和数据
# ------------------------------------------------------------------------------
class TextProcessorApp:
    """
    文本处理应用程序主类
    
    负责创建和管理整个应用程序界面，协调各功能模块。
    
    属性说明:
        root: Tk根窗口
        ai_service: AI服务实例
        template_processor: 模板处理器实例
        ai_extractor: AI文本提取器实例
        current_template: 当前选中的模板路径
        template_fields: 当前模板的字段列表
    
    主要功能:
        - 文本输入和导入
        - AI智能信息提取
        - 模板选择和预览
        - 结果展示和导出
    """
    
    def __init__(self, root: tk.Tk):
        """
        初始化应用程序
        
        参数:
            root: Tk根窗口对象
        
        初始化流程:
            1. 配置窗口属性
            2. 初始化服务实例
            3. 创建菜单栏
            4. 创建主界面
            5. 创建状态栏
            6. 初始化目录结构
        """
        self.root = root
        self.root.title("文本处理与席卡生成程序")
        self.root.geometry("950x650")  # 初始窗口大小
        self.root.minsize(850, 550)  # 最小窗口大小
        
        # 服务实例(延迟初始化)
        self.ai_service: Optional[AIService] = None
        self.template_processor: TemplateProcessor = create_template_processor()
        self.ai_extractor: Optional[AIExtractor] = None
        
        # 当前状态
        self.current_template: str = ""  # 当前模板路径
        
        # 构建界面
        self._create_menu()  # 创建菜单栏
        self._create_main_frame()  # 创建主界面
        self._create_status_bar()  # 创建状态栏
        self._initialize_directories()  # 初始化目录
    
    # --------------------------------------------------------------------------
    # 目录初始化
    # 功能: 确保工作目录存在并刷新模板列表
    # --------------------------------------------------------------------------
    def _initialize_directories(self):
        """
        初始化工作目录
        
        功能说明:
            确保模板目录和输出目录存在，
            并刷新可用模板列表。
        """
        from config import ensure_directories
        ensure_directories(CONFIG)
        self._refresh_template_list()
    
    # --------------------------------------------------------------------------
    # 菜单栏创建
    # 功能: 创建应用程序菜单系统
    # --------------------------------------------------------------------------
    def _create_menu(self):
        """
        创建菜单栏
        
        菜单结构:
            文件菜单:
                - 新建模板
                - 打开模板
                - 导入文本文件
                - 退出
            设置菜单:
                - AI服务设置
                - 输出目录设置
            帮助菜单:
                - 使用说明
                - 关于
        """
        menubar = tk.Menu(self.root)
        
        # 文件菜单
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="新建模板", command=self._new_template)
        file_menu.add_command(label="打开模板", command=self._open_template)
        file_menu.add_separator()  # 分隔线
        file_menu.add_command(label="导入文本文件", command=self._import_text_file)
        file_menu.add_separator()
        file_menu.add_command(label="退出", command=self.root.quit)
        menubar.add_cascade(label="文件", menu=file_menu)
        
        # 设置菜单
        settings_menu = tk.Menu(menubar, tearoff=0)
        settings_menu.add_command(label="AI服务设置", command=self._show_ai_settings)
        settings_menu.add_command(label="输出目录设置", command=self._show_output_settings)
        menubar.add_cascade(label="设置", menu=settings_menu)
        
        # 帮助菜单
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="使用说明", command=self._show_help)
        help_menu.add_command(label="关于", command=self._show_about)
        menubar.add_cascade(label="帮助", menu=help_menu)
        
        self.root.config(menu=menubar)
    
    # --------------------------------------------------------------------------
    # 主界面创建
    # 功能: 创建主工作区域，包括输入区和输出区
    # --------------------------------------------------------------------------
    def _create_main_frame(self):
        """
        创建主界面布局
        
        布局结构:
            左侧 - 输入区域:
                - 文本输入框
                - 席卡生成设置
            右侧 - 输出区域:
                - 模板选择
                - 结果显示区
                - 操作按钮
        """
        # 主框架
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # ==================== 左侧输入区域 ====================
        left_frame = ttk.LabelFrame(main_frame, text="输入区域", padding="10")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        
        # 文本输入标签
        ttk.Label(left_frame, text="请输入人员信息文本:").pack(anchor=tk.W, pady=(0, 5))
        
        # 文本输入框(带滚动条)
        self.input_text = scrolledtext.ScrolledText(left_frame, height=15, wrap=tk.WORD)
        self.input_text.pack(fill=tk.BOTH, expand=True, pady=5)
        
        # 席卡生成框架
        card_settings_frame = ttk.LabelFrame(left_frame, text="席卡生成设置", padding="10")
        card_settings_frame.pack(fill=tk.X, pady=10)
        
        # 活动名称输入
        name_frame = ttk.Frame(card_settings_frame)
        name_frame.pack(fill=tk.X, pady=5)
        ttk.Label(name_frame, text="活动名称:", width=10).pack(side=tk.LEFT, padx=5, anchor=tk.W)
        self.event_name_var = tk.StringVar(value="")
        ttk.Entry(name_frame, textvariable=self.event_name_var, width=40).pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        
        # 席卡显示内容选择
        display_frame = ttk.Frame(card_settings_frame)
        display_frame.pack(fill=tk.X, pady=5)
        ttk.Label(display_frame, text="显示内容:", width=10).pack(side=tk.LEFT, padx=5, anchor=tk.W)
        self.card_display_var = tk.StringVar(value="name")  # 默认显示姓名
        ttk.Radiobutton(display_frame, text="姓名", variable=self.card_display_var, value="name").pack(side=tk.LEFT, padx=10)
        ttk.Radiobutton(display_frame, text="公司名", variable=self.card_display_var, value="company").pack(side=tk.LEFT, padx=10)
        
        # 生成按钮
        button_frame = ttk.Frame(card_settings_frame)
        button_frame.pack(fill=tk.X, pady=5)
        ttk.Button(button_frame, text="生成席卡", command=self._generate_cards, width=20).pack(side=tk.RIGHT, padx=5)
        
        # ==================== 右侧输出区域 ====================
        right_frame = ttk.LabelFrame(main_frame, text="输出区域", padding="10")
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        # 模板选择框架
        template_frame = ttk.LabelFrame(right_frame, text="模板选择", padding="10")
        template_frame.pack(fill=tk.X, pady=10)
        
        # 模板下拉框
        self.template_var = tk.StringVar()
        self.template_combo = ttk.Combobox(template_frame, textvariable=self.template_var, state='readonly')
        self.template_combo.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        self.template_combo.bind('<<ComboboxSelected>>', self._on_template_selected)  # 绑定选择事件
        
        # 浏览按钮
        ttk.Button(template_frame, text="浏览...", command=self._browse_template, width=10).pack(side=tk.RIGHT)
        
        # 结果显示标签
        ttk.Label(right_frame, text="处理结果:").pack(anchor=tk.W, pady=(0, 5))
        
        # 结果显示文本框(只读，带滚动条)
        self.output_text = scrolledtext.ScrolledText(right_frame, height=15, wrap=tk.WORD, state=tk.DISABLED)
        self.output_text.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # 操作按钮框架
        output_button_frame = ttk.Frame(right_frame)
        output_button_frame.pack(fill=tk.X, pady=10)
        
        # 居中对齐按钮
        output_button_frame.columnconfigure(0, weight=1)
        output_button_frame.columnconfigure(1, weight=1)
        
        ttk.Button(output_button_frame, text="打开输出目录", command=self._open_output_dir, width=15).pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)
        ttk.Button(output_button_frame, text="复制结果", command=self._copy_result, width=15).pack(side=tk.LEFT, padx=10, fill=tk.X, expand=True)
    
    # --------------------------------------------------------------------------
    # 状态栏创建
    # 功能: 创建底部状态显示区域
    # --------------------------------------------------------------------------
    def _create_status_bar(self):
        """
        创建状态栏
        
        包含:
            - 状态文本标签
            - 进度条
        """
        # 进度条
        self.progress = ttk.Progressbar(self.root, mode='determinate')
        self.progress.pack(side=tk.BOTTOM, fill=tk.X, pady=(0, 4), padx=10)
        
        # 状态文本
        self.status_var = tk.StringVar(value="就绪")
        status_bar = ttk.Label(self.root, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        status_bar.pack(side=tk.BOTTOM, fill=tk.X, pady=(0, 8), padx=10)
    
    # --------------------------------------------------------------------------
    # 模板列表刷新
    # 功能: 扫描模板目录并更新下拉框列表
    # --------------------------------------------------------------------------
    def _refresh_template_list(self):
        """
        刷新模板列表
        
        功能说明:
            扫描模板目录中的文件，
            筛选出支持的格式(.txt, .docx, .pdf)，
            更新模板下拉框的选项。
        """
        template_dir = CONFIG.template.template_dir
        if os.path.exists(template_dir):
            templates = []
            for f in os.listdir(template_dir):
                if f.endswith(('.txt', '.docx', '.pdf')):
                    templates.append(f)
            self.template_combo['values'] = templates
            
            # 如果有模板且当前未选择，自动选择第一个
            if templates and not self.template_var.get():
                self.template_var.set(templates[0])
                self._on_template_selected(None)
    
    # --------------------------------------------------------------------------
    # 模板选择事件处理
    # 功能: 处理模板下拉框选择事件
    # --------------------------------------------------------------------------
    def _on_template_selected(self, event):
        """
        模板选择事件处理
        
        参数:
            event: 事件对象(可能为None)
        
        功能说明:
            当用户选择模板时更新当前模板路径
        """
        template_name = self.template_var.get()
        if template_name:
            self.current_template = os.path.join(CONFIG.template.template_dir, template_name)
    
    # --------------------------------------------------------------------------
    # 模板浏览
    # 功能: 打开文件选择对话框选择模板
    # --------------------------------------------------------------------------
    def _browse_template(self):
        """
        浏览选择模板文件
        
        功能说明:
            打开文件选择对话框，
            允许用户选择外部的模板文件。
        """
        file_path = filedialog.askopenfilename(
            title="选择模板文件",
            filetypes=[
                ("所有支持格式", "*.txt;*.docx;*.pdf"),
                ("文本文件", "*.txt"),
                ("Word文档", "*.docx"),
                ("PDF文件", "*.pdf")
            ]
        )
        if file_path:
            self.current_template = file_path
            self.template_var.set(os.path.basename(file_path))
            self._on_template_selected(None)
    
    # --------------------------------------------------------------------------
    # 错误显示
    # 功能: 在输出区域显示错误信息
    # --------------------------------------------------------------------------
    def _display_error(self, error_message: str):
        """
        显示错误信息
        
        参数:
            error_message: 错误信息文本
        
        功能说明:
            在输出区域显示错误信息，
            同时更新状态栏和进度条。
        """
        self.progress['value'] = 0
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(tk.END, f"错误: {error_message}\n")
        self.output_text.config(state=tk.DISABLED)
        self.status_var.set("处理失败")
    
    # --------------------------------------------------------------------------
    # AI服务检查
    # 功能: 检查并初始化AI服务
    # --------------------------------------------------------------------------
    def _check_ai_service(self) -> bool:
        """
        检查AI服务是否可用
        
        返回值:
            bool: AI服务是否可用
        
        功能说明:
            1. 检查AI服务实例是否存在
            2. 检查API密钥是否配置
            3. 如果未配置，弹出设置对话框
        """
        if self.ai_service is None:
            if not CONFIG.ai.api_key:
                self._show_ai_settings()
                return False
            self.ai_service = create_ai_service()
        
        if not self.ai_service.config.api_key:
            messagebox.showwarning("警告", "请先配置AI服务API密钥")
            self._show_ai_settings()
            return False
        
        return True
    
    # --------------------------------------------------------------------------
    # AI提取器检查
    # 功能: 检查并初始化AI提取器
    # --------------------------------------------------------------------------
    def _check_ai_extractor(self) -> bool:
        """
        检查AI提取器是否可用
        
        返回值:
            bool: AI提取器是否可用
        
        功能说明:
            1. 先检查AI服务
            2. 初始化或更新AI提取器实例
        """
        if not self._check_ai_service():
            return False
        
        if self.ai_extractor is None:
            self.ai_extractor = create_extractor(self.ai_service)
        else:
            self.ai_extractor.set_ai_service(self.ai_service)
        
        return True
    

    

    
    # --------------------------------------------------------------------------
    # 席卡生成功能
    # 功能: 根据提取的人员信息生成席卡文档
    # --------------------------------------------------------------------------
    def _generate_cards(self):
        """
        生成席卡文档
        
        功能说明:
            1. 获取输入文本并提取人员信息
            2. 获取活动名称和显示内容设置
            3. 使用席卡模板生成docx文档
            4. 保存到输出目录
        
        使用场景:
            用于会议、活动等场合的席卡打印，
            支持选择显示姓名或公司名。
        """
        text = self.input_text.get(1.0, tk.END).strip()
        if not text:
            messagebox.showwarning("警告", "请输入文本内容")
            return
        
        if not self._check_ai_extractor():
            return
        
        # 获取设置
        event_name = self.event_name_var.get().strip()
        display_type = self.card_display_var.get()  # "name" 或 "company"
        
        # 检查是否选择了席卡模板
        if not self.current_template:
            messagebox.showwarning("警告", "请先选择席卡模板")
            return
        
        # 检查模板文件类型
        template_ext = os.path.splitext(self.current_template)[1].lower()
        if template_ext not in ['.docx', '.doc']:
            messagebox.showwarning("警告", "席卡生成功能仅支持.docx格式的模板文件，请选择一个Word文档模板")
            return
        
        self.status_var.set("正在生成席卡...")
        self.progress['value'] = 0
        
        def generate_thread():
            try:
                self.root.after(0, lambda: self.progress.config(value=20))
                
                # 提取人员信息
                infos = self.ai_extractor.extract_from_text(text)
                
                if not infos:
                    self.root.after(0, lambda: messagebox.showwarning("警告", "未能提取到人员信息"))
                    return
                
                self.root.after(0, lambda: self.progress.config(value=50))
                
                # 生成席卡
                result = self._create_card_document(infos, event_name, display_type)
                
                self.root.after(0, lambda: self.progress.config(value=100))
                
                if result['success']:
                    # 构建详细的结果显示
                    valid_count = len([f for f in result['files'] if f.get('valid', False)])
                    invalid_count = len([f for f in result['files'] if not f.get('valid', False)])
                    
                    display_text = f"席卡生成完成\n"
                    display_text += f"=" * 40 + "\n"
                    display_text += f"输出目录: {result['output_dir']}\n"
                    display_text += f"成功生成: {valid_count} 个文件\n"
                    if invalid_count > 0:
                        display_text += f"验证失败: {invalid_count} 个文件\n"
                    display_text += f"\n已生成文件:\n"
                    display_text += "-" * 40 + "\n"
                    
                    for f in result['files']:
                        status = "✓" if f.get('valid', False) else "✗"
                        display_text += f"{status} {f['filename']}\n"
                    
                    if result.get('failed'):
                        display_text += f"\n失败记录: {', '.join(result['failed'])}\n"
                    
                    display_text += f"\n质量报告: {result['report_path']}"
                    
                    # 直接更新输出文本
                    self.root.after(0, lambda: self._update_output_text(display_text, "席卡生成结果"))
                    self.root.after(0, lambda: self.status_var.set(f"席卡生成完成，共 {valid_count} 个有效文件"))
                else:
                    self.root.after(0, lambda: self._display_error(result['error']))
                    
            except Exception as e:
                self.root.after(0, lambda: self._display_error(str(e)))
        
        threading.Thread(target=generate_thread, daemon=True).start()
    
    # --------------------------------------------------------------------------
    # 创建席卡文档
    # 功能: 实际生成席卡docx文档（单人单文件模式）
    # --------------------------------------------------------------------------
    def _create_card_document(self, infos, event_name: str, display_type: str) -> dict:
        """
        创建席卡文档（单人单文件模式）
        
        参数:
            infos: 人员信息列表(PersonInfo对象)
            event_name: 活动名称
            display_type: 显示类型("name"或"company")
        
        返回值:
            dict: 包含success、count、output_dir、files、errors的字典
        
        实现说明:
            采用单人单文件模式，每个人生成一个独立的docx文件。
            创建独立文件夹存储所有席卡，包含时间戳和任务标识。
            实现空白页检测和质量验证机制。
            支持Word转PDF和PDF合并功能。
        """
        import logging
        from datetime import datetime
        
        # 配置日志
        log_data = {
            'start_time': datetime.now().isoformat(),
            'event_name': event_name,
            'display_type': display_type,
            'total_persons': len(infos),
            'errors': [],
            'warnings': [],
            'generated_files': []
        }
        
        try:
            from docx import Document
            
            # 加载模板
            template_path = self.current_template
            if not os.path.exists(template_path):
                return {'success': False, 'error': f'模板文件不存在: {template_path}', 'log': log_data}
            
            # 创建独立文件夹：席卡_活动名称_时间戳
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            event_suffix = f"_{event_name}" if event_name else ""
            folder_name = f"席卡{event_suffix}_{timestamp}"
            output_dir = os.path.join(CONFIG.template.output_dir, folder_name)
            os.makedirs(output_dir, exist_ok=True)
            
            # 创建word和pdf子文件夹
            word_dir = os.path.join(output_dir, "word")
            pdf_dir = os.path.join(output_dir, "pdf")
            os.makedirs(word_dir, exist_ok=True)
            os.makedirs(pdf_dir, exist_ok=True)
            
            log_data['output_dir'] = output_dir
            log_data['word_dir'] = word_dir
            log_data['pdf_dir'] = pdf_dir
            
            # 记录生成结果
            generated_files = []
            failed_persons = []
            card_count = 0
            pdf_files = []
            
            # 遍历人员信息，为每个人生成一个独立文件
            for i, info in enumerate(infos):
                # 确定显示内容
                if display_type == "company":
                    display_text = info.company if info.company else info.name
                else:
                    display_text = info.name
                
                if not display_text:
                    log_data['warnings'].append(f"第{i+1}条记录无有效显示内容，已跳过")
                    continue
                
                try:
                    # 为每个人重新加载模板
                    doc = Document(template_path)
                    
                    # 替换文档中的占位符
                    for paragraph in doc.paragraphs:
                        self._replace_placeholder_in_paragraph(paragraph, display_text, event_name, display_type, info)
                    
                    # 替换表格中的占位符
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    self._replace_placeholder_in_paragraph(paragraph, display_text, event_name, display_type, info)
                    
                    # 生成文件名：席卡_姓名_日期
                    date_str = datetime.now().strftime("%Y%m%d")
                    safe_name = self._sanitize_filename(display_text)
                    filename = f"席卡_{safe_name}_{date_str}.docx"
                    filepath = os.path.join(word_dir, filename)
                    
                    # 处理文件名冲突
                    counter = 1
                    while os.path.exists(filepath):
                        filename = f"席卡_{safe_name}_{date_str}_{counter}.docx"
                        filepath = os.path.join(word_dir, filename)
                        counter += 1
                    
                    # 保存文档
                    doc.save(filepath)
                    
                    # 质量验证：检测空白页
                    # 注意：两字姓名已被格式化（插入全角空格），验证时需要检查原始姓名
                    validation_result = self._validate_card_document(filepath, display_text, display_type)
                    
                    if validation_result['valid']:
                        # 尝试将Word转换为PDF
                        pdf_filename = filename.replace('.docx', '.pdf')
                        pdf_filepath = os.path.join(pdf_dir, pdf_filename)
                        
                        if self._convert_docx_to_pdf(filepath, pdf_filepath):
                            pdf_files.append(pdf_filepath)
                            generated_files.append({
                                'filename': filename,
                                'filepath': filepath,
                                'pdf_filepath': pdf_filepath,
                                'name': display_text,
                                'valid': True
                            })
                            card_count += 1
                        else:
                            log_data['warnings'].append(f"文件 {filename} 转换为PDF失败")
                            generated_files.append({
                                'filename': filename,
                                'filepath': filepath,
                                'name': display_text,
                                'valid': True,
                                'warning': 'PDF转换失败'
                            })
                            card_count += 1
                    else:
                        # 空白或内容不完整，尝试重新生成
                        log_data['warnings'].append(f"文件 {filename} 验证失败: {validation_result['reason']}")
                        generated_files.append({
                            'filename': filename,
                            'filepath': filepath,
                            'name': display_text,
                            'valid': False,
                            'reason': validation_result['reason']
                        })
                        failed_persons.append(display_text)
                    
                except Exception as e:
                    error_msg = f"生成 {display_text} 的席卡失败: {str(e)}"
                    log_data['errors'].append(error_msg)
                    failed_persons.append(display_text)
            
            # 生成PDF合集
            pdf_combined_path = None
            if pdf_files:
                pdf_combined_path = os.path.join(output_dir, f"席卡合集_{timestamp}.pdf")
                if self._merge_pdfs(pdf_files, pdf_combined_path):
                    log_data['pdf_combined_path'] = pdf_combined_path
                else:
                    log_data['warnings'].append("PDF合集生成失败")
            
            # 生成质量报告
            report_path = os.path.join(output_dir, "生成报告.txt")
            self._generate_quality_report(report_path, log_data, generated_files, failed_persons)
            
            log_data['end_time'] = datetime.now().isoformat()
            log_data['generated_count'] = card_count
            log_data['failed_count'] = len(failed_persons)
            log_data['pdf_count'] = len(pdf_files)
            
            return {
                'success': True,
                'count': card_count,
                'output_dir': output_dir,
                'word_dir': word_dir,
                'pdf_dir': pdf_dir,
                'files': generated_files,
                'failed': failed_persons,
                'log': log_data,
                'report_path': report_path,
                'pdf_combined_path': pdf_combined_path
            }
            
        except ImportError:
            return {'success': False, 'error': '请先安装python-docx库: pip install python-docx', 'log': log_data}
        except Exception as e:
            import traceback
            log_data['errors'].append(str(e))
            log_data['traceback'] = traceback.format_exc()
            return {'success': False, 'error': str(e), 'log': log_data}
    
    # --------------------------------------------------------------------------
    # 文件名安全处理
    # 功能: 移除文件名中的非法字符
    # --------------------------------------------------------------------------
    def _sanitize_filename(self, name: str) -> str:
        """
        处理文件名，移除非法字符
        
        参数:
            name: 原始名称
        
        返回值:
            str: 安全的文件名
        """
        import re
        # 移除Windows文件名中的非法字符
        illegal_chars = r'[<>:"/\\|?*]'
        safe_name = re.sub(illegal_chars, '_', name)
        # 限制文件名长度
        return safe_name[:50] if len(safe_name) > 50 else safe_name
    
    # --------------------------------------------------------------------------
    # 席卡文档质量验证
    # 功能: 检测文档是否为空白或内容不完整
    # --------------------------------------------------------------------------
    def _validate_card_document(self, filepath: str, expected_content: str, display_type: str = "name") -> dict:
        """
        验证席卡文档质量
        
        参数:
            filepath: 文档路径
            expected_content: 期望包含的内容
            display_type: 显示类型("name"或"company")
        
        返回值:
            dict: 包含valid和reason的字典
        """
        try:
            from docx import Document
            
            doc = Document(filepath)
            
            # 提取所有文本内容
            all_text = []
            for paragraph in doc.paragraphs:
                all_text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text.append(cell.text)
            
            full_text = ''.join(all_text)
            
            # 检查是否为空白
            if not full_text.strip():
                return {'valid': False, 'reason': '文档内容为空'}
            
            # 检查文档结构完整性
            if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
                return {'valid': False, 'reason': '文档无段落或表格'}
            
            # 检查是否包含期望的内容
            # 注意：两字姓名会被格式化为"姓　名"（中间插入全角空格）
            if expected_content:
                # 先检查原始内容
                if expected_content in full_text:
                    return {'valid': True, 'reason': '验证通过'}
                
                # 如果是姓名且为两字，检查格式化后的内容
                if display_type == "name" and len(expected_content) == 2:
                    formatted_content = expected_content[0] + '\u3000' + expected_content[1]
                    if formatted_content in full_text:
                        return {'valid': True, 'reason': '验证通过'}
                
                # 检查是否包含去除空格后的内容（兼容处理）
                normalized_text = full_text.replace('\u3000', '').replace(' ', '')
                if expected_content in normalized_text:
                    return {'valid': True, 'reason': '验证通过'}
                
                return {'valid': False, 'reason': f'文档未包含期望内容: {expected_content}'}
            
            return {'valid': True, 'reason': '验证通过'}
            
        except Exception as e:
            return {'valid': False, 'reason': f'验证过程出错: {str(e)}'}
    
    def _convert_docx_to_pdf(self, docx_path: str, pdf_path: str) -> bool:
        """
        将Word文档转换为PDF
        
        参数:
            docx_path: Word文档路径
            pdf_path: 输出PDF路径
        
        返回值:
            bool: 转换是否成功
        """
        try:
            import win32com.client
            import os
            
            # 使用Word.Application进行转换
            word = win32com.client.Dispatch('Word.Application')
            word.Visible = False
            
            try:
                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17是PDF格式
                doc.Close()
                return True
            finally:
                word.Quit()
        except Exception as e:
            # 转换失败，尝试使用spire.doc作为备选方案
            try:
                from spire.doc import Document as SpireDocument
                
                doc = SpireDocument()
                doc.LoadFromFile(docx_path)
                doc.SaveToFile(pdf_path, 1)
                doc.Close()
                return True
            except:
                return False
    
    def _merge_pdfs(self, pdf_files: list, output_path: str) -> bool:
        """
        合并多个PDF文件为一个
        
        参数:
            pdf_files: PDF文件路径列表
            output_path: 输出文件路径
        
        返回值:
            bool: 合并是否成功
        """
        try:
            from PyPDF2 import PdfWriter, PdfReader
            
            merger = PdfWriter()
            
            for pdf_file in pdf_files:
                with open(pdf_file, 'rb') as f:
                    reader = PdfReader(f)
                    for page_num in range(len(reader.pages)):
                        page = reader.pages[page_num]
                        merger.add_page(page)
            
            with open(output_path, 'wb') as f:
                merger.write(f)
            
            return True
        except Exception as e:
            return False
    
    # --------------------------------------------------------------------------
    # 生成质量报告
    # 功能: 生成详细的席卡生成质量报告
    # --------------------------------------------------------------------------
    def _generate_quality_report(self, report_path: str, log_data: dict, generated_files: list, failed_persons: list):
        """
        生成质量报告文件
        
        参数:
            report_path: 报告文件路径
            log_data: 日志数据
            generated_files: 已生成的文件列表
            failed_persons: 失败的人员列表
        """
        from datetime import datetime
        import os
        
        report_lines = [
            "=" * 60,
            "席卡生成质量报告".center(50),
            "=" * 60,
            "",
            f"生成时间: {log_data.get('start_time', 'N/A')}",
            f"活动名称: {log_data.get('event_name', '未指定')}",
            f"显示类型: {'公司名' if log_data.get('display_type') == 'company' else '姓名'}",
            f"总人数: {log_data.get('total_persons', 0)}",
            f"成功生成: {len([f for f in generated_files if f.get('valid', False)])}",
            f"验证失败: {len([f for f in generated_files if not f.get('valid', False)])}",
            f"PDF转换: {log_data.get('pdf_count', 0)}",
            "",
            f"输出目录: {log_data.get('output_dir', 'N/A')}",
            f"Word目录: {log_data.get('word_dir', 'N/A')}",
            f"PDF目录: {log_data.get('pdf_dir', 'N/A')}",
            "",
        ]
        
        if log_data.get('pdf_combined_path'):
            report_lines.extend([
                f"PDF合集: {os.path.basename(log_data['pdf_combined_path'])}",
                "",
            ])
        
        report_lines.extend([
            "-" * 60,
            "已生成文件列表:",
            "-" * 60,
        ])
        
        for f in generated_files:
            status = "✓ 有效" if f.get('valid', False) else f"✗ {f.get('reason', '无效')}"
            pdf_status = " (已转PDF)" if f.get('pdf_filepath') else ""
            report_lines.append(f"  {f['filename']}{pdf_status} - {status}")
        
        if failed_persons:
            report_lines.extend([
                "",
                "-" * 60,
                "失败记录:",
                "-" * 60,
            ])
            for name in failed_persons:
                report_lines.append(f"  • {name}")
        
        if log_data.get('warnings'):
            report_lines.extend([
                "",
                "-" * 60,
                "警告信息:",
                "-" * 60,
            ])
            for warning in log_data['warnings']:
                report_lines.append(f"  ⚠ {warning}")
        
        if log_data.get('errors'):
            report_lines.extend([
                "",
                "-" * 60,
                "错误信息:",
                "-" * 60,
            ])
            for error in log_data['errors']:
                report_lines.append(f"  ✗ {error}")
        
        report_lines.extend([
            "",
            "=" * 60,
            f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            "=" * 60,
        ])
        
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(report_lines))
    
    # --------------------------------------------------------------------------
    # 替换段落中的占位符
    # 功能: 在docx段落中替换席卡相关占位符并设置字体格式
    # --------------------------------------------------------------------------
    def _replace_placeholder_in_paragraph(self, paragraph, display_text: str, event_name: str, display_type: str, info):
        """
        替换段落中的占位符并设置字体格式
        
        参数:
            paragraph: docx段落对象
            display_text: 要显示的文本(姓名或公司名)
            event_name: 活动名称
            display_type: 显示类型("name"或"company")
            info: 人员信息对象
        
        占位符说明:
            {{name}} - 姓名
            {{company}} - 公司名
            {{position}} - 职位
            {{event}} - 活动名称
            {{display}} - 显示内容(根据用户选择)
            {姓名/公司} - 姓名或公司名(根据用户选择)
            {活动名称} - 活动名称
        
        字体格式:
            活动名称: 楷体, 三号(16pt)
            姓名: 楷体, 72pt (两字姓名中间插入全角空格)
            公司名称: 楷体, 36pt
        
        注意:
            字体名称必须是Windows系统中已安装的字体。
            常见中文字体: 楷体(KaiTi)、宋体(SimSun)、黑体(SimHei)、微软雅黑(Microsoft YaHei)
        """
        from datetime import datetime
        from docx.shared import Pt
        from docx.oxml.ns import qn
        
        # 格式化姓名：两字姓名中间插入全角空格
        def format_name(text):
            """
            格式化姓名文本
            规则：如果姓名由两个汉字组成，在中间插入全角空格(U+3000)
            """
            if not text:
                return text
            # 检查是否为两个汉字
            if len(text) == 2 and all('\u4e00' <= c <= '\u9fff' for c in text):
                return text[0] + '\u3000' + text[1]  # 插入全角空格
            return text
        
        # 判断是否为公司名称（用于确定字号）
        # 根据display_type判断：company表示公司名，name表示姓名
        is_company = (display_type == "company")
        
        # 根据类型确定字号：姓名72pt，公司名36pt
        name_font_size = 36 if is_company else 72
        
        # 格式化显示文本
        formatted_display = display_text if is_company else format_name(display_text)
        formatted_name = info.name if is_company else format_name(info.name)
        
        # 定义替换映射及其字体设置
        # 格式: (占位符, 替换值, 字体名称, 字号)
        # 注意: 字体名称必须是Windows系统中已安装的字体
        placeholder_configs = [
            # 活动名称: 楷体, 三号(16pt)
            ('{活动名称}', event_name, '楷体', 16),
            ('{{event}}', event_name, '楷体', 16),
            # 姓名/公司: 楷体，字号根据类型确定
            ('{姓名/公司}', formatted_display, '楷体', name_font_size),
            ('{{display}}', formatted_display, '楷体', name_font_size),
            ('{{name}}', formatted_name, '楷体', name_font_size),
            ('{{company}}', info.company, '楷体', 36),  # 公司名称固定36pt
            # 其他字段: 使用默认字体
            ('{{position}}', info.position, None, None),
            ('{{date}}', datetime.now().strftime("%Y-%m-%d"), None, None),
        ]
        
        # 合并段落中所有run的文本
        full_text = ''.join([run.text for run in paragraph.runs])
        
        # 检查是否包含任何占位符
        found_placeholders = []
        for placeholder, value, font_name, font_size in placeholder_configs:
            if placeholder in full_text:
                found_placeholders.append((placeholder, value, font_name, font_size))
        
        if not found_placeholders:
            return
        
        # 清空段落中所有run
        for run in paragraph.runs:
            run.text = ""
        
        # 处理每个找到的占位符
        remaining_text = full_text
        first_run = paragraph.runs[0] if paragraph.runs else paragraph.add_run("")
        
        for placeholder, value, font_name, font_size in found_placeholders:
            if placeholder in remaining_text:
                # 分割文本
                parts = remaining_text.split(placeholder, 1)
                
                # 添加占位符前的文本（保持原格式）
                if parts[0]:
                    if paragraph.runs:
                        current_run = paragraph.runs[-1]
                        current_run.text = current_run.text + parts[0]
                    else:
                        paragraph.add_run(parts[0])
                
                # 添加替换后的文本（应用新格式）
                if value:
                    new_run = paragraph.add_run(str(value))
                    
                    # 设置字体
                    if font_name:
                        try:
                            new_run.font.name = font_name
                            # 设置中文字体
                            r = new_run._element
                            rPr = r.get_or_add_rPr()
                            rFonts = rPr.get_or_add_rFonts()
                            rFonts.set(qn('w:eastAsia'), font_name)
                        except Exception:
                            pass  # 字体设置失败时忽略
                    
                    # 设置字号
                    if font_size:
                        try:
                            new_run.font.size = Pt(font_size)
                        except Exception:
                            pass
                
                # 更新剩余文本
                remaining_text = parts[1] if len(parts) > 1 else ""
        
        # 添加剩余文本
        if remaining_text:
            if paragraph.runs:
                current_run = paragraph.runs[-1]
                current_run.text = current_run.text + remaining_text
            else:
                paragraph.add_run(remaining_text)
    

    

    
    # --------------------------------------------------------------------------
    # 更新输出文本
    # 功能: 在输出区域显示结果文本
    # --------------------------------------------------------------------------
    def _update_output_text(self, result: str, title: str):
        """
        更新输出文本
        
        参数:
            result: 结果文本
            title: 结果标题
        """
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete(1.0, tk.END)
        self.output_text.insert(tk.END, f"{title}:\n")
        self.output_text.insert(tk.END, "-" * 40 + "\n")
        self.output_text.insert(tk.END, result)
        self.output_text.config(state=tk.DISABLED)
    
    # --------------------------------------------------------------------------
    # 复制结果
    # 功能: 将结果复制到剪贴板
    # --------------------------------------------------------------------------
    def _copy_result(self):
        """
        复制结果到剪贴板
        
        功能说明:
            将输出区域的内容复制到系统剪贴板。
        """
        content = self.output_text.get(1.0, tk.END)
        if not content.strip():
            messagebox.showwarning("警告", "没有可复制的内容")
            return
        
        self.root.clipboard_clear()
        self.root.clipboard_append(content)
        self.status_var.set("结果已复制到剪贴板")
    
    # --------------------------------------------------------------------------
    # 保存结果
    # 功能: 将结果保存到文件
    # --------------------------------------------------------------------------
    def _save_result(self):
        """
        保存结果到文件
        
        功能说明:
            打开文件保存对话框，
            将输出区域的内容保存为文本文件。
        """
        content = self.output_text.get(1.0, tk.END)
        if not content.strip():
            messagebox.showwarning("警告", "没有可保存的内容")
            return
        
        file_path = filedialog.asksaveasfilename(
            title="保存结果",
            defaultextension=".txt",
            filetypes=[
                ("文本文件", "*.txt"),
                ("所有文件", "*.*")
            ]
        )
        
        if file_path:
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(content)
            messagebox.showinfo("成功", f"结果已保存到: {file_path}")
    
    # --------------------------------------------------------------------------
    # 打开输出目录
    # 功能: 在文件管理器中打开输出目录
    # --------------------------------------------------------------------------
    def _open_output_dir(self):
        """
        打开输出目录
        
        功能说明:
            在系统文件管理器中打开输出目录。
        """
        output_dir = CONFIG.template.output_dir
        if os.path.exists(output_dir):
            os.startfile(output_dir)
    
    # --------------------------------------------------------------------------
    # 新建模板
    # 功能: 打开新建模板对话框
    # --------------------------------------------------------------------------
    def _new_template(self):
        """
        新建模板
        
        功能说明:
            打开模板创建对话框，
            创建完成后刷新模板列表。
        """
        dialog = TemplateDialog(self.root)
        self.root.wait_window(dialog.top)
        self._refresh_template_list()
    
    # --------------------------------------------------------------------------
    # 打开模板
    # 功能: 浏览选择模板文件
    # --------------------------------------------------------------------------
    def _open_template(self):
        """打开模板文件(调用浏览功能)"""
        self._browse_template()
    
    # --------------------------------------------------------------------------
    # 导入文本文件
    # 功能: 从文件导入文本到输入区域
    # --------------------------------------------------------------------------
    def _import_text_file(self):
        """
        导入文本文件
        
        功能说明:
            打开文件选择对话框，
            将选中的文本文件内容导入到输入区域。
        """
        file_path = filedialog.askopenfilename(
            title="选择文本文件",
            filetypes=[("文本文件", "*.txt"), ("所有文件", "*.*")]
        )
        if file_path:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            self.input_text.delete(1.0, tk.END)
            self.input_text.insert(tk.END, content)
    
    # --------------------------------------------------------------------------
    # AI设置对话框
    # 功能: 显示AI服务配置对话框
    # --------------------------------------------------------------------------
    def _show_ai_settings(self):
        """
        显示AI服务设置对话框
        
        功能说明:
            打开AI配置对话框，
            用户可以修改API密钥、地址、模型等设置。
        """
        dialog = AISettingsDialog(self.root, CONFIG.ai)
        self.root.wait_window(dialog.top)
        
        if dialog.result:
            CONFIG.ai = dialog.result
            self.ai_service = create_ai_service(CONFIG.ai)
    
    # --------------------------------------------------------------------------
    # 输出设置对话框
    # 功能: 显示输出目录配置对话框
    # --------------------------------------------------------------------------
    def _show_output_settings(self):
        """
        显示输出设置对话框
        
        功能说明:
            打开输出配置对话框，
            用户可以修改模板目录和输出目录。
        """
        dialog = OutputSettingsDialog(self.root, CONFIG.template)
        self.root.wait_window(dialog.top)
        
        if dialog.result:
            CONFIG.template = dialog.result
            self.template_processor = create_template_processor()
    
    # --------------------------------------------------------------------------
    # 帮助对话框
    # 功能: 显示使用说明
    # --------------------------------------------------------------------------
    def _show_help(self):
        """
        显示使用说明
        
        功能说明:
            在新窗口中显示详细的使用说明文档。
        """
        help_text = """
自动化文本处理与模板填充程序 - 使用说明

1. 配置AI服务
   - 在"设置"菜单中选择"AI服务设置"
   - 输入API密钥和API地址

2. 选择模板
   - 从下拉列表选择已有模板
   - 或点击"浏览"选择其他模板文件

3. 输入文本
   - 在左侧输入区域输入或粘贴文本
   - 或从文件导入文本

4. AI智能提取
   - 勾选需要提取的内容（姓名/单位/职位）
   - 点击"提取"、"表格"或"打印"按钮
   - 程序将使用AI智能提取信息

5. 席卡生成功能
   - 输入活动名称（可选）
   - 选择席卡显示内容：姓名或公司名
   - 选择席卡模板文件（.docx格式）
   - 点击"生成席卡"按钮
   - 程序将自动生成席卡文档

模板格式说明:
- 使用 {{字段名}} 作为占位符
- 例如: {{name}}, {{date}}, {{content}}
- 席卡模板支持: {{name}}, {{company}}, {{position}}, {{event}}, {{display}}, {{date}}
- {{display}} 会根据用户选择显示姓名或公司名
- 程序会自动替换占位符为实际值
        """
        
        help_window = tk.Toplevel(self.root)
        help_window.title("使用说明")
        help_window.geometry("500x400")
        
        text = scrolledtext.ScrolledText(help_window, wrap=tk.WORD)
        text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        text.insert(tk.END, help_text)
        text.config(state=tk.DISABLED)
    
    # --------------------------------------------------------------------------
    # 关于对话框
    # 功能: 显示程序信息
    # --------------------------------------------------------------------------
    def _show_about(self):
        """显示关于对话框"""
        messagebox.showinfo(
            "关于",
            "自动化文本处理与模板填充程序\n\n"
            "版本: 1.0.0\n\n"
            "功能:\n"
            "- AI智能文本处理\n"
            "- 模板自动填充\n"
            "- 多种格式输出"
        )


# ------------------------------------------------------------------------------
# AI设置对话框类
# 功能: 提供AI服务配置界面
# ------------------------------------------------------------------------------
class AISettingsDialog:
    """
    AI服务设置对话框
    
    属性说明:
        result: 对话框结果(确认时为AIConfig对象)
        config: 当前配置
        top: 对话框窗口
    
    配置项:
        - API密钥
        - API地址
        - 模型名称
        - 最大Tokens
        - Temperature
        - 超时时间
    """
    
    def __init__(self, parent, config: AIConfig):
        """
        初始化对话框
        
        参数:
            parent: 父窗口
            config: 当前AI配置
        """
        self.result = None
        self.config = config
        
        # 创建对话框窗口
        self.top = tk.Toplevel(parent)
        self.top.title("AI服务设置")
        self.top.geometry("450x300")
        self.top.transient(parent)  # 设置为父窗口的临时窗口
        self.top.grab_set()  # 模态对话框
        
        # 创建表单
        frame = ttk.Frame(self.top, padding="10")
        frame.pack(fill=tk.BOTH, expand=True)
        
        # API密钥
        ttk.Label(frame, text="API密钥:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.api_key_var = tk.StringVar(value=config.api_key)
        self.api_key_entry = ttk.Entry(frame, textvariable=self.api_key_var, width=40, show="*")  # 密码模式
        self.api_key_entry.grid(row=0, column=1, pady=5, padx=5)
        
        # API地址
        ttk.Label(frame, text="API地址:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.api_url_var = tk.StringVar(value=config.api_base_url)
        ttk.Entry(frame, textvariable=self.api_url_var, width=40).grid(row=1, column=1, pady=5, padx=5)
        
        # 模型
        ttk.Label(frame, text="模型:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.model_var = tk.StringVar(value=config.model)
        ttk.Entry(frame, textvariable=self.model_var, width=40).grid(row=2, column=1, pady=5, padx=5)
        
        # 最大Tokens
        ttk.Label(frame, text="最大Tokens:").grid(row=3, column=0, sticky=tk.W, pady=5)
        self.max_tokens_var = tk.StringVar(value=str(config.max_tokens))
        ttk.Entry(frame, textvariable=self.max_tokens_var, width=40).grid(row=3, column=1, pady=5, padx=5)
        
        # Temperature
        ttk.Label(frame, text="Temperature:").grid(row=4, column=0, sticky=tk.W, pady=5)
        self.temperature_var = tk.StringVar(value=str(config.temperature))
        ttk.Entry(frame, textvariable=self.temperature_var, width=40).grid(row=4, column=1, pady=5, padx=5)
        
        # 超时
        ttk.Label(frame, text="超时(秒):").grid(row=5, column=0, sticky=tk.W, pady=5)
        self.timeout_var = tk.StringVar(value=str(config.timeout))
        ttk.Entry(frame, textvariable=self.timeout_var, width=40).grid(row=5, column=1, pady=5, padx=5)
        
        # 按钮
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=6, column=0, columnspan=2, pady=20)
        
        ttk.Button(button_frame, text="确定", command=self._ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=self._cancel).pack(side=tk.LEFT, padx=5)
    
    def _ok(self):
        """
        确认按钮处理
        
        功能说明:
            验证输入并创建AIConfig对象，
            成功后关闭对话框。
        """
        try:
            self.result = AIConfig(
                api_key=self.api_key_var.get(),
                api_base_url=self.api_url_var.get(),
                model=self.model_var.get(),
                max_tokens=int(self.max_tokens_var.get()),
                temperature=float(self.temperature_var.get()),
                timeout=int(self.timeout_var.get())
            )
            self.top.destroy()
        except ValueError as e:
            messagebox.showerror("错误", f"输入值无效: {e}")
    
    def _cancel(self):
        """取消按钮处理，关闭对话框"""
        self.top.destroy()


# ------------------------------------------------------------------------------
# 输出设置对话框类
# 功能: 提供输出目录配置界面
# ------------------------------------------------------------------------------
class OutputSettingsDialog:
    """
    输出设置对话框
    
    属性说明:
        result: 对话框结果(确认时为TemplateConfig对象)
        config: 当前配置
        top: 对话框窗口
    
    配置项:
        - 模板目录
        - 输出目录
    """
    
    def __init__(self, parent, config):
        """
        初始化对话框
        
        参数:
            parent: 父窗口
            config: 当前模板配置
        """
        self.result = None
        self.config = config
        
        # 创建对话框窗口
        self.top = tk.Toplevel(parent)
        self.top.title("输出设置")
        self.top.geometry("450x150")
        self.top.transient(parent)
        self.top.grab_set()
        
        # 创建表单
        frame = ttk.Frame(self.top, padding="10")
        frame.pack(fill=tk.BOTH, expand=True)
        
        # 模板目录
        ttk.Label(frame, text="模板目录:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.template_dir_var = tk.StringVar(value=config.template_dir)
        ttk.Entry(frame, textvariable=self.template_dir_var, width=30).grid(row=0, column=1, pady=5, padx=5)
        ttk.Button(frame, text="浏览...", command=lambda: self._browse_dir(self.template_dir_var)).grid(row=0, column=2)
        
        # 输出目录
        ttk.Label(frame, text="输出目录:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.output_dir_var = tk.StringVar(value=config.output_dir)
        ttk.Entry(frame, textvariable=self.output_dir_var, width=30).grid(row=1, column=1, pady=5, padx=5)
        ttk.Button(frame, text="浏览...", command=lambda: self._browse_dir(self.output_dir_var)).grid(row=1, column=2)
        
        # 按钮
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=2, column=0, columnspan=3, pady=20)
        
        ttk.Button(button_frame, text="确定", command=self._ok).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=self._cancel).pack(side=tk.LEFT, padx=5)
    
    def _browse_dir(self, var: tk.StringVar):
        """
        浏览目录
        
        参数:
            var: 存储目录路径的StringVar对象
        """
        dir_path = filedialog.askdirectory()
        if dir_path:
            var.set(dir_path)
    
    def _ok(self):
        """
        确认按钮处理
        
        功能说明:
            创建TemplateConfig对象并关闭对话框。
        """
        from config import TemplateConfig
        self.result = TemplateConfig(
            template_dir=self.template_dir_var.get(),
            output_dir=self.output_dir_var.get()
        )
        self.top.destroy()
    
    def _cancel(self):
        """取消按钮处理"""
        self.top.destroy()


# ------------------------------------------------------------------------------
# 模板创建对话框类
# 功能: 提供新建模板的界面
# ------------------------------------------------------------------------------
class TemplateDialog:
    """
    模板创建对话框
    
    属性说明:
        result: 对话框结果
        top: 对话框窗口
    
    功能:
        - 输入模板名称
        - 编辑模板内容
        - 创建模板文件
    """
    
    def __init__(self, parent):
        """
        初始化对话框
        
        参数:
            parent: 父窗口
        """
        self.result = None
        
        # 创建对话框窗口
        self.top = tk.Toplevel(parent)
        self.top.title("新建模板")
        self.top.geometry("500x400")
        self.top.transient(parent)
        self.top.grab_set()
        
        # 创建表单
        frame = ttk.Frame(self.top, padding="10")
        frame.pack(fill=tk.BOTH, expand=True)
        
        # 模板名称
        ttk.Label(frame, text="模板名称:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.name_var = tk.StringVar(value="new_template.txt")
        ttk.Entry(frame, textvariable=self.name_var, width=40).grid(row=0, column=1, pady=5)
        
        # 模板内容
        ttk.Label(frame, text="模板内容:").grid(row=1, column=0, sticky=tk.NW, pady=5)
        self.content_text = scrolledtext.ScrolledText(frame, width=50, height=15)
        self.content_text.grid(row=1, column=1, pady=5)
        
        # 插入示例模板内容
        self.content_text.insert(tk.END, """====================================
              文档模板
====================================

标题: {{title}}
日期: {{date}}

尊敬的 {{name}}:

您好！

{{content}}

此致
敬礼！

====================================
""")
        
        # 按钮
        button_frame = ttk.Frame(frame)
        button_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        ttk.Button(button_frame, text="创建", command=self._create).pack(side=tk.LEFT, padx=5)
        ttk.Button(button_frame, text="取消", command=self.top.destroy).pack(side=tk.LEFT, padx=5)
    
    def _create(self):
        """
        创建模板文件
        
        功能说明:
            验证输入并创建模板文件。
        """
        name = self.name_var.get().strip()
        content = self.content_text.get(1.0, tk.END)
        
        if not name:
            messagebox.showwarning("警告", "请输入模板名称")
            return
        
        template_path = os.path.join(CONFIG.template.template_dir, name)
        
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        messagebox.showinfo("成功", f"模板已创建: {template_path}")
        self.top.destroy()


# ------------------------------------------------------------------------------
# GUI启动函数
# 功能: 创建并运行GUI应用程序
# ------------------------------------------------------------------------------
def run_gui():
    """
    启动GUI应用程序
    
    功能说明:
        创建Tk根窗口和应用程序实例，
        启动主事件循环。
    """
    root = tk.Tk()
    app = TextProcessorApp(root)
    root.mainloop()


# 模块入口
if __name__ == "__main__":
    run_gui()
