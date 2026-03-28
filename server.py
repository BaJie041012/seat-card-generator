# ==============================================================================
# 文件名称: server.py
# 功能描述: 网络服务模块，提供HTTP接口供同一网段内的其他设备访问
# 创建日期: 2026-03-25
# 作    者: 戒者有八
# 版本: 1.0.0
# 描述: 基于Flask的HTTP服务，提供席卡生成的Web界面和API接口
# ==============================================================================
"""
网络服务模块 - 基于Flask的HTTP服务

本模块提供网络服务功能，允许同一网段内的其他设备通过HTTP访问应用程序。
功能包括:
    - 提供Web界面访问应用功能
    - 支持席卡生成的API接口
    - 配置网络监听地址和端口
    - 支持Word转PDF和PDF合并功能

设计思路:
    1. 使用Flask框架构建轻量级HTTP服务
    2. 提供RESTful API接口供客户端调用
    3. 集成AI服务进行人员信息提取
    4. 支持文件下载功能

使用方式:
    1. 启动服务: python start_server.py
    2. 访问Web界面: http://localhost:5000
    3. 调用API接口: POST /api/generate-cards
"""

# 导入必要的模块
import os  # 用于文件和目录操作
import sys  # 用于系统操作
import json  # 用于JSON数据处理
from datetime import datetime  # 用于处理日期时间
from typing import Dict, Any, Optional  # 用于类型注解
from flask import Flask, request, render_template_string, jsonify, send_from_directory  # Flask框架

# 添加项目根目录到Python路径，确保能导入项目模块
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

# 导入项目模块
from config import CONFIG, ensure_directories  # 配置管理
from ai_service import create_ai_service  # AI服务创建
from template_processor import create_template_processor  # 模板处理器
from text_extractor import create_extractor  # 文本提取器

# 初始化Flask应用实例
app = Flask(__name__)

# 服务配置
SERVER_HOST = '0.0.0.0'  # 监听所有网络接口，允许局域网内其他设备访问
SERVER_PORT = 5000  # 默认端口，可根据需要修改

# 增加超时时间，适应AI处理和文件操作
REQUEST_TIMEOUT = 60  # 60秒超时，避免长时间操作被中断

# 文件下载映射（用于避免路径编码问题）
# 键为文件ID，值为文件路径
file_download_map = {}

# 初始化服务实例
template_processor = create_template_processor()  # 创建模板处理器实例
ai_service = None  # AI服务实例，延迟初始化
extractor = None  # 文本提取器实例，延迟初始化

# 确保目录存在，避免文件操作时目录不存在的错误
ensure_directories(CONFIG)

# ------------------------------------------------------------------------------
# 辅助函数
# ------------------------------------------------------------------------------
def get_local_ip() -> str:
    """
    获取本机局域网IP地址
    
    返回值:
        str: 本机局域网IP地址
    
    实现说明:
        通过创建UDP套接字连接到外部服务器(8.8.8.8:80)来获取本地IP地址
        如果获取失败，返回本地回环地址127.0.0.1
    """
    import socket  # 导入socket模块用于网络操作
    try:
        # 创建套接字连接到外部服务器以获取本地IP
        # 使用UDP协议，不需要实际发送数据
        s = socket.socket(socket.AF_INET, socket.SOCK_DGRAM)
        # 连接到Google DNS服务器(8.8.8.8)的80端口
        s.connect(('8.8.8.8', 80))
        # 获取套接字的本地地址，提取IP部分
        ip = s.getsockname()[0]
        # 关闭套接字
        s.close()
        # 返回获取到的IP地址
        return ip
    except Exception:
        # 如果获取失败，返回本地回环地址
        return '127.0.0.1'

# ------------------------------------------------------------------------------
# Web界面路由
# ------------------------------------------------------------------------------
@app.route('/')
def index():
    """
    主页面路由
    
    功能: 渲染Web界面，提供席卡生成的用户界面
    
    返回值:
        str: 渲染后的HTML页面
    
    实现说明:
        1. 获取本机局域网IP地址
        2. 使用render_template_string渲染HTML模板
        3. 传入local_ip和port参数，用于在页面中显示服务地址
    """
    # 获取本机局域网IP地址，用于在页面中显示
    local_ip = get_local_ip()
    return render_template_string('''
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>文本处理与模板填充服务</title>
        <style>
            body {
                font-family: Arial, sans-serif;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                line-height: 1.6;
            }
            h1 {
                color: #333;
                text-align: center;
            }
            h2 {
                color: #555;
                margin-top: 30px;
            }
            .info-box {
                background-color: #f5f5f5;
                padding: 15px;
                border-radius: 5px;
                margin: 20px 0;
            }
            .code {
                background-color: #e0e0e0;
                padding: 10px;
                border-radius: 3px;
                font-family: monospace;
                white-space: pre-wrap;
            }
            .endpoint {
                margin: 15px 0;
                padding: 10px;
                border-left: 3px solid #4CAF50;
                background-color: #f9f9f9;
            }
            .button {
                display: inline-block;
                padding: 10px 20px;
                background-color: #4CAF50;
                color: white;
                text-decoration: none;
                border-radius: 4px;
                margin: 5px;
                border: none;
                cursor: pointer;
                font-size: 16px;
            }
            .button:hover {
                background-color: #45a049;
            }
            .form-group {
                margin: 15px 0;
            }
            label {
                display: block;
                margin-bottom: 5px;
                font-weight: bold;
            }
            input[type="text"], select, textarea {
                width: 100%;
                padding: 10px;
                border: 1px solid #ddd;
                border-radius: 4px;
                box-sizing: border-box;
            }
            textarea {
                height: 150px;
                resize: vertical;
            }
            .form-container {
                background-color: #f9f9f9;
                padding: 20px;
                border-radius: 5px;
                margin: 20px 0;
            }
            .result {
                background-color: #e8f5e8;
                padding: 15px;
                border-radius: 5px;
                margin: 20px 0;
                border-left: 4px solid #4CAF50;
            }
            .error {
                background-color: #ffebee;
                padding: 15px;
                border-radius: 5px;
                margin: 20px 0;
                border-left: 4px solid #f44336;
            }
            .loading {
                display: inline-block;
                width: 20px;
                height: 20px;
                border: 3px solid rgba(0,0,0,.3);
                border-radius: 50%;
                border-top-color: #4CAF50;
                animation: spin 1s ease-in-out infinite;
            }
            @keyframes spin {
                to { transform: rotate(360deg); }
            }
            .hidden {
                display: none;
            }
        </style>
    </head>
    <body>
        <h1>文本处理与模板填充服务</h1>
        
        <div class="info-box">
            <h2>服务信息</h2>
            <p><strong>服务地址:</strong> http://{{ local_ip }}:{{ port }}</p>
            <p><strong>局域网访问:</strong> 同一网段内的设备可通过上述地址访问</p>
        </div>
        
        <h2>席卡生成</h2>
        <div class="form-container">
            <form id="cardForm">
                <div class="form-group">
                    <label for="text">人员信息文本</label>
                    <textarea id="text" name="text" placeholder="请输入人员信息，例如：张三，公司A，职位B；李四，公司C，职位D"></textarea>
                </div>
                <div class="form-group">
                    <label for="event_name">活动名称</label>
                    <input type="text" id="event_name" name="event_name" placeholder="例如：公司年会">
                </div>
                <div class="form-group">
                    <label for="display_type">显示内容</label>
                    <select id="display_type" name="display_type">
                        <option value="name">姓名</option>
                        <option value="company">公司名</option>
                    </select>
                </div>
                <div class="form-group">
                    <label for="template">模板文件</label>
                    <select id="template" name="template"></select>
                </div>
                <div style="text-align: center;">
                    <button type="submit" class="button">生成席卡</button>
                    <div id="loading" class="loading hidden"></div>
                </div>
            </form>
        </div>
        
        <div id="result" class="hidden"></div>
        
        <h2>使用指南</h2>
        <ol>
            <li>确保防火墙已允许端口 {{ port }} 的入站连接</li>
            <li>在同一网段的设备上，打开浏览器访问: <code>http://{{ local_ip }}:{{ port }}</code></li>
            <li>在上方的表单中输入人员信息，选择模板，点击生成席卡按钮</li>
            <li>等待生成完成后，查看生成结果和输出文件</li>
        </ol>
        
        <div class="info-box">
            <h3>注意事项</h3>
            <ul>
                <li>服务默认运行在端口 {{ port }}</li>
                <li>确保AI服务API密钥已正确配置</li>
                <li>模板文件应放在 templates 目录中</li>
                <li>生成的文件将保存在 output 目录中</li>
            </ul>
        </div>
        
        <div style="margin-top: 30px; text-align: center;">
            <a href="/api/templates" class="button">查看模板列表</a>
            <a href="/api/output" class="button">查看输出目录</a>
        </div>
        
        <script>
            // 加载模板列表
            window.onload = function() {
                fetch('/api/templates')
                    .then(response => response.json())
                    .then(data => {
                        if (data.success) {
                            const templateSelect = document.getElementById('template');
                            data.templates.forEach(template => {
                                const option = document.createElement('option');
                                option.value = template;
                                option.textContent = template;
                                templateSelect.appendChild(option);
                            });
                        }
                    });
            };
            
            // 表单提交
            document.getElementById('cardForm').addEventListener('submit', function(e) {
                e.preventDefault();
                
                const formData = {
                    text: document.getElementById('text').value,
                    event_name: document.getElementById('event_name').value,
                    display_type: document.getElementById('display_type').value,
                    template: document.getElementById('template').value
                };
                
                // 显示加载状态
                document.getElementById('loading').classList.remove('hidden');
                document.getElementById('result').classList.add('hidden');
                
                // 发送请求
                fetch('/api/generate-cards', {
                    method: 'POST',
                    headers: {
                        'Content-Type': 'application/json'
                    },
                    body: JSON.stringify(formData)
                })
                .then(response => response.json())
                .then(data => {
                    // 隐藏加载状态
                    document.getElementById('loading').classList.add('hidden');
                    
                    // 显示结果
                    const resultDiv = document.getElementById('result');
                    resultDiv.classList.remove('hidden');
                    
                    if (data.success) {
                        resultDiv.className = 'result';
                        resultDiv.innerHTML = `
                            <h3>生成成功!</h3>
                            <p><strong>生成数量:</strong> ${data.count}</p>
                            <p><strong>输出目录:</strong> ${data.output_dir}</p>
                            <p><strong>Word目录:</strong> ${data.word_dir}</p>
                            <p><strong>PDF目录:</strong> ${data.pdf_dir}</p>
                            
                            <p><strong>Word文件:</strong></p>
                            <ul>
                                ${data.word_files.map(file => {
                                    return `<li><a href="/download/${file.file_id}" target="_blank">${file.filename}</a></li>`;
                                }).join('')}
                            </ul>
                            
                            ${data.pdf_files.length > 0 ? `
                                <p><strong>PDF文件:</strong></p>
                                <ul>
                                    ${data.pdf_files.map(file => {
                                        return `<li><a href="/download/${file.file_id}" target="_blank">${file.filename}</a></li>`;
                                    }).join('')}
                                </ul>
                            ` : ''}
                            
                            ${data.pdf_combined_id ? `
                                <p><strong>PDF合集:</strong> <a href="/download/${data.pdf_combined_id}" target="_blank">席卡合集.pdf</a></p>
                            ` : ''}
                            
                            ${data.failed.length > 0 ? `<p><strong>失败记录:</strong> ${data.failed.join(', ')}</p>` : ''}
                            <p><strong>质量报告:</strong> 
                                ${data.report_id ? 
                                    `<a href="/download/${data.report_id}" target="_blank">查看报告</a>` : 
                                    '无报告'
                                }
                            </p>
                        `;
                    } else {
                        resultDiv.className = 'error';
                        resultDiv.innerHTML = `
                            <h3>生成失败</h3>
                            <p>${data.error}</p>
                        `;
                    }
                })
                .catch(error => {
                    // 隐藏加载状态
                    document.getElementById('loading').classList.add('hidden');
                    
                    // 显示错误
                    const resultDiv = document.getElementById('result');
                    resultDiv.className = 'error';
                    resultDiv.classList.remove('hidden');
                    resultDiv.innerHTML = `
                        <h3>请求出错</h3>
                        <p>请检查网络连接或服务状态</p>
                    `;
                });
            });
        </script>
    </body>
    </html>
    ''', local_ip=local_ip, port=SERVER_PORT)

# ------------------------------------------------------------------------------
# API路由
# ------------------------------------------------------------------------------
@app.route('/api/templates', methods=['GET'])
def get_templates():
    """
    获取可用模板列表
    """
    template_dir = CONFIG.template.template_dir
    templates = []
    
    if os.path.exists(template_dir):
        for f in os.listdir(template_dir):
            if f.endswith(('.txt', '.docx', '.pdf')):
                templates.append(f)
    
    return jsonify({
        'success': True,
        'templates': templates
    })

@app.route('/api/output', methods=['GET'])
def get_output():
    """
    获取输出目录内容
    """
    output_dir = CONFIG.template.output_dir
    files = []
    
    if os.path.exists(output_dir):
        for root, dirs, filenames in os.walk(output_dir):
            # 计算相对路径
            rel_path = os.path.relpath(root, output_dir)
            if rel_path == '.':
                rel_path = ''
            
            for filename in filenames:
                files.append({
                    'path': os.path.join(rel_path, filename),
                    'size': os.path.getsize(os.path.join(root, filename))
                })
    
    return jsonify({
        'success': True,
        'output_dir': output_dir,
        'files': files
    })

@app.route('/download/<file_id>')
def download_file(file_id):
    """
    下载文件 - 使用文件ID
    """
    # 检查文件ID是否存在
    if file_id not in file_download_map:
        return jsonify({
            'success': False,
            'error': '文件不存在或已过期'
        }), 404
    
    # 获取文件路径
    full_path = file_download_map[file_id]
    
    # 检查文件是否存在
    if not os.path.isfile(full_path):
        # 从映射中移除不存在的文件
        del file_download_map[file_id]
        return jsonify({
            'success': False,
            'error': '文件不存在'
        }), 404
    
    # 使用send_file发送文件
    from flask import send_file
    return send_file(full_path, as_attachment=True)

@app.route('/api/generate-cards', methods=['POST'])
def generate_cards():
    """
    生成席卡
    """
    global ai_service, extractor
    
    # 获取请求数据
    data = request.json
    
    if not data:
        return jsonify({
            'success': False,
            'error': '请提供请求数据'
        }), 400
    
    text = data.get('text')
    event_name = data.get('event_name', '')
    display_type = data.get('display_type', 'name')
    template_name = data.get('template', '席卡模板v4.docx')
    
    if not text:
        return jsonify({
            'success': False,
            'error': '请提供人员信息文本'
        }), 400
    
    # 初始化AI服务
    if ai_service is None:
        if not CONFIG.ai.api_key:
            return jsonify({
                'success': False,
                'error': 'AI服务API密钥未配置'
            }), 500
        # 创建AI服务实例并设置更长的超时时间
        ai_config = CONFIG.ai
        ai_config.timeout = REQUEST_TIMEOUT
        ai_service = create_ai_service(ai_config)
    
    # 初始化提取器
    if extractor is None:
        extractor = create_extractor(ai_service)
    
    try:
        # 提取人员信息
        infos = extractor.extract_from_text(text)
        
        if not infos:
            return jsonify({
                'success': False,
                'error': '未能提取到人员信息'
            }), 400
        
        # 加载模板
        template_path = os.path.join(CONFIG.template.template_dir, template_name)
        if not os.path.exists(template_path):
            return jsonify({
                'success': False,
                'error': f'模板文件不存在: {template_name}'
            }), 404
        
        # 检查模板文件类型
        template_ext = os.path.splitext(template_path)[1].lower()
        if template_ext not in ['.docx', '.doc']:
            return jsonify({
                'success': False,
                'error': '席卡生成功能仅支持.docx格式的模板文件'
            }), 400
        
        # 生成席卡
        import docx
        
        # 创建独立文件夹：席卡_活动名称_时间戳
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        event_suffix = f"_{event_name}" if event_name else ""
        folder_name = f"席卡{event_suffix}_{timestamp}"
        output_dir = os.path.join(CONFIG.template.output_dir, folder_name)
        os.makedirs(output_dir, exist_ok=True)
        
        # 创建word和pdf子文件夹
        word_dir = os.path.join(output_dir, "word")
        pdf_dir = os.path.join(output_dir, "pdf")
        os.makedirs(word_dir, exist_ok=True)
        os.makedirs(pdf_dir, exist_ok=True)
        
        # 记录生成结果
        generated_files = []
        failed_persons = []
        card_count = 0
        pdf_files = []
        
        # 遍历人员信息，为每个人生成一个独立文件
        for i, info in enumerate(infos):
            # 确定显示内容
            if display_type == "company":
                display_text = info.company if info.company else info.name
            else:
                display_text = info.name
            
            if not display_text:
                failed_persons.append(f"第{i+1}条记录")
                continue
            
            try:
                # 为每个人重新加载模板
                doc = docx.Document(template_path)
                
                # 替换文档中的占位符
                for paragraph in doc.paragraphs:
                    _replace_placeholder_in_paragraph(paragraph, display_text, event_name, display_type, info)
                
                # 替换表格中的占位符
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                _replace_placeholder_in_paragraph(paragraph, display_text, event_name, display_type, info)
                
                # 生成文件名：席卡_姓名_日期
                date_str = datetime.now().strftime("%Y%m%d")
                safe_name = _sanitize_filename(display_text)
                filename = f"席卡_{safe_name}_{date_str}.docx"
                filepath = os.path.join(word_dir, filename)
                
                # 处理文件名冲突
                counter = 1
                while os.path.exists(filepath):
                    filename = f"席卡_{safe_name}_{date_str}_{counter}.docx"
                    filepath = os.path.join(word_dir, filename)
                    counter += 1
                
                # 保存文档
                doc.save(filepath)
                
                # 尝试将Word转换为PDF
                pdf_filename = filename.replace('.docx', '.pdf')
                pdf_filepath = os.path.join(pdf_dir, pdf_filename)
                
                if _convert_docx_to_pdf(filepath, pdf_filepath):
                    pdf_files.append(pdf_filepath)
                    generated_files.append({
                        'filename': filename,
                        'filepath': filepath,
                        'pdf_filepath': pdf_filepath,
                        'name': display_text
                    })
                else:
                    generated_files.append({
                        'filename': filename,
                        'filepath': filepath,
                        'name': display_text
                    })
                card_count += 1
                
            except Exception as e:
                failed_persons.append(display_text)
        
        # 生成PDF合集
        pdf_combined_path = None
        if pdf_files:
            pdf_combined_path = os.path.join(output_dir, f"席卡合集_{timestamp}.pdf")
            if _merge_pdfs(pdf_files, pdf_combined_path):
                pass
        
        # 生成质量报告
        report_path = os.path.join(output_dir, "生成报告.txt")
        _generate_quality_report(report_path, event_name, display_type, len(infos), generated_files, failed_persons, pdf_files, pdf_combined_path, word_dir, pdf_dir)
        
        # 为每个文件生成唯一ID，用于下载
        import uuid
        download_files = []
        pdf_files_info = []
        for file in generated_files:
            # 为Word文件生成下载ID
            file_id = str(uuid.uuid4())
            file_path = file['filepath']
            file_download_map[file_id] = file_path
            
            file_info = {
                'filename': file['filename'],
                'file_id': file_id,
                'type': 'word'
            }
            
            # 为PDF文件生成下载ID
            if file.get('pdf_filepath'):
                pdf_file_id = str(uuid.uuid4())
                pdf_file_path = file['pdf_filepath']
                file_download_map[pdf_file_id] = pdf_file_path
                
                file_info['pdf_filename'] = os.path.basename(pdf_file_path)
                file_info['pdf_file_id'] = pdf_file_id
                
                pdf_files_info.append({
                    'filename': os.path.basename(pdf_file_path),
                    'file_id': pdf_file_id
                })
            
            download_files.append(file_info)
        
        # 为报告生成下载ID
        report_id = None
        if report_path:
            report_id = str(uuid.uuid4())
            file_download_map[report_id] = report_path
        
        # 为PDF合集生成下载ID
        pdf_combined_id = None
        if pdf_combined_path:
            pdf_combined_id = str(uuid.uuid4())
            file_download_map[pdf_combined_id] = pdf_combined_path
        
        return jsonify({
            'success': True,
            'count': card_count,
            'output_dir': output_dir,
            'word_files': download_files,
            'pdf_files': pdf_files_info,
            'failed': failed_persons,
            'report_id': report_id,
            'pdf_combined_id': pdf_combined_id,
            'word_dir': word_dir,
            'pdf_dir': pdf_dir
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# ------------------------------------------------------------------------------
# 辅助函数
# ------------------------------------------------------------------------------
def _sanitize_filename(name: str) -> str:
    """
    处理文件名，移除非法字符
    """
    import re
    # 移除Windows文件名中的非法字符
    illegal_chars = r'[<>:"/\\|?*]'
    safe_name = re.sub(illegal_chars, '_', name)
    # 限制文件名长度
    return safe_name[:50] if len(safe_name) > 50 else safe_name

def _replace_placeholder_in_paragraph(paragraph, display_text: str, event_name: str, display_type: str, info):
    """
    替换段落中的占位符并设置字体格式
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    
    # 格式化姓名：两字姓名中间插入全角空格
    def format_name(text):
        if not text:
            return text
        # 检查是否为两个汉字
        if len(text) == 2 and all('\u4e00' <= c <= '\u9fff' for c in text):
            return text[0] + '\u3000' + text[1]  # 插入全角空格
        return text
    
    # 判断是否为公司名称（用于确定字号）
    is_company = (display_type == "company")
    
    # 根据类型确定字号：姓名72pt，公司名36pt
    name_font_size = 36 if is_company else 72
    
    # 格式化显示文本
    formatted_display = display_text if is_company else format_name(display_text)
    formatted_name = info.name if is_company else format_name(info.name)
    
    # 定义替换映射及其字体设置
    # 格式: (占位符, 替换值, 字体名称, 字号)
    placeholder_configs = [
        # 活动名称: 楷体, 三号(16pt)
        ('{活动名称}', event_name, '楷体', 16),
        ('{{event}}', event_name, '楷体', 16),
        # 姓名/公司: 楷体，字号根据类型确定
        ('{姓名/公司}', formatted_display, '楷体', name_font_size),
        ('{{display}}', formatted_display, '楷体', name_font_size),
        ('{{name}}', formatted_name, '楷体', name_font_size),
        ('{{company}}', info.company, '楷体', 36),  # 公司名称固定36pt
        # 其他字段: 使用默认字体
        ('{{position}}', info.position, None, None),
        ('{{date}}', datetime.now().strftime("%Y-%m-%d"), None, None),
    ]
    
    # 合并段落中所有run的文本
    full_text = ''.join([run.text for run in paragraph.runs])
    
    # 检查是否包含任何占位符
    found_placeholders = []
    for placeholder, value, font_name, font_size in placeholder_configs:
        if placeholder in full_text:
            found_placeholders.append((placeholder, value, font_name, font_size))
    
    if not found_placeholders:
        return
    
    # 清空段落中所有run
    for run in paragraph.runs:
        run.text = ""
    
    # 处理每个找到的占位符
    remaining_text = full_text
    
    for placeholder, value, font_name, font_size in found_placeholders:
        if placeholder in remaining_text:
            # 分割文本
            parts = remaining_text.split(placeholder, 1)
            
            # 添加占位符前的文本（保持原格式）
            if parts[0]:
                if paragraph.runs:
                    current_run = paragraph.runs[-1]
                    current_run.text = current_run.text + parts[0]
                else:
                    paragraph.add_run(parts[0])
            
            # 添加替换后的文本（应用新格式）
            if value:
                new_run = paragraph.add_run(str(value))
                
                # 设置字体
                if font_name:
                    try:
                        new_run.font.name = font_name
                        # 设置中文字体
                        r = new_run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), font_name)
                    except Exception:
                        pass  # 字体设置失败时忽略
                
                # 设置字号
                if font_size:
                    try:
                        new_run.font.size = Pt(font_size)
                    except Exception:
                        pass
            
            # 更新剩余文本
            remaining_text = parts[1] if len(parts) > 1 else ""
    
    # 添加剩余文本
    if remaining_text:
        if paragraph.runs:
            current_run = paragraph.runs[-1]
            current_run.text = current_run.text + remaining_text
        else:
            paragraph.add_run(remaining_text)

def _generate_quality_report(report_path: str, event_name: str, display_type: str, total_persons: int, generated_files: list, failed_persons: list, pdf_files: list = None, pdf_combined_path: str = None, word_dir: str = None, pdf_dir: str = None):
    """
    生成质量报告文件
    """
    import os
    
    report_lines = [
        "=" * 60,
        "席卡生成质量报告".center(50),
        "=" * 60,
        "",
        f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"活动名称: {event_name or '未指定'}",
        f"显示类型: {'公司名' if display_type == 'company' else '姓名'}",
        f"总人数: {total_persons}",
        f"成功生成: {len(generated_files)}",
        f"生成失败: {len(failed_persons)}",
        f"PDF转换: {len(pdf_files) if pdf_files else 0}",
        "",
        f"Word目录: {word_dir or 'N/A'}",
        f"PDF目录: {pdf_dir or 'N/A'}",
        "",
    ]
    
    if pdf_combined_path:
        report_lines.extend([
            f"PDF合集: {os.path.basename(pdf_combined_path)}",
            "",
        ])
    
    report_lines.extend([
        "-" * 60,
        "已生成文件列表:",
        "-" * 60,
    ])
    
    for f in generated_files:
        pdf_status = " (已转PDF)" if f.get('pdf_filepath') else ""
        report_lines.append(f"  • {f['filename']}{pdf_status}")
    
    if failed_persons:
        report_lines.extend([
            "",
            "-" * 60,
            "失败记录:",
            "-" * 60,
        ])
        for name in failed_persons:
            report_lines.append(f"  • {name}")
    
    report_lines.extend([
        "",
        "=" * 60,
        f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 60,
    ])
    
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report_lines))

def _convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """
    将Word文档转换为PDF
    """
    try:
        import win32com.client
        import os
        
        # 使用Word.Application进行转换
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        
        try:
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17是PDF格式
            doc.Close()
            return True
        finally:
            word.Quit()
    except Exception as e:
        # 转换失败，尝试使用spire.doc作为备选方案
        try:
            from spire.doc import Document as SpireDocument
            
            doc = SpireDocument()
            doc.LoadFromFile(docx_path)
            doc.SaveToFile(pdf_path, 1)
            doc.Close()
            return True
        except:
            return False

def _merge_pdfs(pdf_files: list, output_path: str) -> bool:
    """
    合并多个PDF文件为一个
    """
    try:
        from PyPDF2 import PdfWriter, PdfReader
        
        merger = PdfWriter()
        
        for pdf_file in pdf_files:
            with open(pdf_file, 'rb') as f:
                reader = PdfReader(f)
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    merger.add_page(page)
        
        with open(output_path, 'wb') as f:
            merger.write(f)
        
        return True
    except Exception as e:
        return False

# ------------------------------------------------------------------------------
# 启动服务
# ------------------------------------------------------------------------------
def run_server(host=SERVER_HOST, port=SERVER_PORT):
    """
    启动Flask服务
    
    参数:
        host: 监听地址
        port: 监听端口
    """
    print(f"\n=== 文本处理与模板填充服务 ===")
    print(f"服务启动中...")
    print(f"本地访问地址: http://localhost:{port}")
    print(f"局域网访问地址: http://{get_local_ip()}:{port}")
    print(f"\n按 Ctrl+C 停止服务\n")
    
    app.run(host=host, port=port, debug=False)

if __name__ == '__main__':
    run_server()
