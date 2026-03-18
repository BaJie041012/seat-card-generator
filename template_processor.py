# ==============================================================================
# 文件名称: template_processor.py
# 功能描述: 模板处理模块，支持docx、pdf、txt格式的模板填充和输出
# 创建日期: 2026-03-14
# 作    者: 戒者有八
# ==============================================================================
"""
模板处理模块 - 支持docx、pdf、txt格式的模板填充

本模块提供模板文件的处理功能，包括:
    - 加载不同格式的模板文件
    - 识别模板中的占位符
    - 使用数据填充模板
    - 保存填充后的文件

设计思路:
    采用策略模式，为每种文件格式定义独立的处理器类。
    通过抽象基类定义统一接口，实现多态调用。
    支持扩展新的文件格式，只需添加新的处理器类。
"""

import os
import re
import json
from typing import Dict, Any, Optional, List
from dataclasses import dataclass
from datetime import datetime
from abc import ABC, abstractmethod


# ------------------------------------------------------------------------------
# 模板处理结果数据类
# 功能: 封装模板处理的结果信息
# ------------------------------------------------------------------------------
@dataclass
class TemplateResult:
    """
    模板处理结果数据类
    
    属性说明:
        success: 处理是否成功
        output_path: 输出文件路径，成功时有效
        error_message: 错误信息，失败时有效
    
    设计思路:
        统一成功和失败的返回格式，便于调用方判断和处理。
    """
    success: bool  # 处理成功标志
    output_path: str = ""  # 输出文件路径
    error_message: str = ""  # 错误信息


# ------------------------------------------------------------------------------
# 模板处理器抽象基类
# 功能: 定义模板处理器的统一接口
# 设计思路: 使用抽象基类强制子类实现核心方法，保证接口一致性
# ------------------------------------------------------------------------------
class BaseTemplateProcessor(ABC):
    """
    模板处理器抽象基类
    
    定义所有模板处理器必须实现的方法:
        - load_template: 加载模板文件
        - fill_template: 填充模板数据
        - save: 保存处理结果
        - get_placeholders: 获取占位符列表(可选)
    
    设计模式:
        模板方法模式，定义算法骨架，子类实现具体步骤。
    """
    
    @abstractmethod
    def load_template(self, template_path: str) -> bool:
        """
        加载模板文件
        
        参数:
            template_path: 模板文件路径
        
        返回值:
            bool: 加载是否成功
        """
        pass
    
    @abstractmethod
    def fill_template(self, data: Dict[str, Any]) -> str:
        """
        使用数据填充模板
        
        参数:
            data: 键值对数据，键为占位符名称，值为替换内容
        
        返回值:
            str: 填充后的内容或状态信息
        """
        pass
    
    @abstractmethod
    def save(self, output_path: str) -> bool:
        """
        保存处理结果到文件
        
        参数:
            output_path: 输出文件路径
        
        返回值:
            bool: 保存是否成功
        """
        pass
    
    def get_placeholders(self) -> List[str]:
        """
        获取模板中的占位符列表
        
        返回值:
            List[str]: 占位符名称列表
        
        默认实现:
            返回空列表，子类可重写以提供实际功能。
        """
        return []


# ------------------------------------------------------------------------------
# 文本模板处理器
# 功能: 处理纯文本格式的模板文件
# 设计思路: 最基础的处理器，直接进行字符串替换
# ------------------------------------------------------------------------------
class TxtTemplateProcessor(BaseTemplateProcessor):
    """
    文本模板处理器
    
    处理.txt格式的模板文件，支持简单的占位符替换。
    
    属性说明:
        template_content: 原始模板内容
        filled_content: 填充后的内容
    
    占位符格式:
        使用双花括号格式: {{字段名}}
        例如: {{name}}, {{date}}, {{content}}
    """
    
    def __init__(self):
        """初始化文本模板处理器"""
        self.template_content = ""  # 原始模板内容
        self.filled_content = ""  # 填充后的内容
    
    def load_template(self, template_path: str) -> bool:
        """
        加载文本模板文件
        
        参数:
            template_path: 模板文件路径
        
        返回值:
            bool: 加载是否成功
        
        实现说明:
            以UTF-8编码读取文件内容，存储到template_content属性。
        """
        try:
            with open(template_path, 'r', encoding='utf-8') as f:
                self.template_content = f.read()
            return True
        except Exception as e:
            return False
    
    def fill_template(self, data: Dict[str, Any]) -> str:
        """
        填充文本模板
        
        参数:
            data: 键值对数据字典
        
        返回值:
            str: 填充后的文本内容
        
        实现说明:
            1. 复制模板内容到filled_content
            2. 遍历数据字典，替换对应的占位符
            3. 自动填充date和time占位符
        """
        self.filled_content = self.template_content
        
        # 替换数据字段
        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            self.filled_content = self.filled_content.replace(placeholder, str(value) if value else "")
        
        # 自动填充日期和时间
        self.filled_content = self.filled_content.replace("{{date}}", datetime.now().strftime("%Y-%m-%d"))
        self.filled_content = self.filled_content.replace("{{time}}", datetime.now().strftime("%H:%M:%S"))
        
        return self.filled_content
    
    def save(self, output_path: str) -> bool:
        """
        保存填充结果到文件
        
        参数:
            output_path: 输出文件路径
        
        返回值:
            bool: 保存是否成功
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(self.filled_content)
            return True
        except Exception as e:
            return False
    
    def get_placeholders(self) -> List[str]:
        """
        提取模板中的占位符
        
        返回值:
            List[str]: 占位符名称列表
        
        实现说明:
            使用正则表达式匹配{{字段名}}格式，
            提取其中的字段名称。
        """
        pattern = r'\{\{(\w+)\}\}'  # 匹配{{字段名}}格式
        return re.findall(pattern, self.template_content)


# ------------------------------------------------------------------------------
# Word文档模板处理器
# 功能: 处理.docx格式的Word文档模板
# 设计思路: 使用python-docx库操作Word文档，支持段落和表格
# ------------------------------------------------------------------------------
class DocxTemplateProcessor(BaseTemplateProcessor):
    """
    Word文档模板处理器
    
    处理.docx格式的Word文档模板，支持段落和表格中的占位符替换。
    
    属性说明:
        document: python-docx文档对象
        template_path: 模板文件路径
    
    依赖:
        需要安装python-docx库: pip install python-docx
    
    特点:
        - 支持段落中的占位符替换
        - 支持表格中的占位符替换
        - 保持原有格式不变
    """
    
    def __init__(self):
        """初始化Word文档模板处理器"""
        self.document = None  # 文档对象
        self.template_path = ""  # 模板路径
    
    def load_template(self, template_path: str) -> bool:
        """
        加载Word文档模板
        
        参数:
            template_path: 模板文件路径
        
        返回值:
            bool: 加载是否成功
        
        异常处理:
            - ImportError: python-docx库未安装
            - 其他异常: 文件读取失败
        """
        try:
            from docx import Document
            self.document = Document(template_path)
            self.template_path = template_path
            return True
        except ImportError:
            return False  # 库未安装
        except Exception as e:
            return False  # 其他错误
    
    def fill_template(self, data: Dict[str, Any]) -> str:
        """
        填充Word文档模板
        
        参数:
            data: 键值对数据字典
        
        返回值:
            str: 处理状态信息
        
        实现说明:
            1. 复制数据并添加日期时间字段
            2. 遍历所有段落进行替换
            3. 遍历所有表格单元格进行替换
        """
        if self.document is None:
            return ""
        
        # 准备数据，添加日期时间
        data_with_datetime = data.copy()
        data_with_datetime['date'] = datetime.now().strftime("%Y-%m-%d")
        data_with_datetime['time'] = datetime.now().strftime("%H:%M:%S")
        
        # 替换段落中的占位符
        for paragraph in self.document.paragraphs:
            self._replace_in_paragraph(paragraph, data_with_datetime)
        
        # 替换表格中的占位符
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data_with_datetime)
        
        return "Document filled successfully"
    
    def _replace_in_paragraph(self, paragraph, data: Dict[str, Any]):
        """
        在段落中替换占位符
        
        参数:
            paragraph: docx段落对象
            data: 数据字典
        
        实现说明:
            遍历段落中的所有run(文本片段)，
            在每个run中查找并替换占位符。
            使用run级别替换可以保持原有格式。
        """
        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value) if value else "")
    
    def save(self, output_path: str) -> bool:
        """
        保存Word文档
        
        参数:
            output_path: 输出文件路径
        
        返回值:
            bool: 保存是否成功
        """
        if self.document is None:
            return False
        try:
            self.document.save(output_path)
            return True
        except Exception as e:
            return False
    
    def get_placeholders(self) -> List[str]:
        """
        提取Word文档中的占位符
        
        返回值:
            List[str]: 占位符名称列表(去重)
        
        实现说明:
            遍历文档中的所有段落和表格单元格，
            使用正则表达式提取占位符。
        """
        if self.document is None:
            return []
        
        placeholders = set()  # 使用集合去重
        pattern = r'\{\{(\w+)\}\}'
        
        # 提取段落中的占位符
        for paragraph in self.document.paragraphs:
            matches = re.findall(pattern, paragraph.text)
            placeholders.update(matches)
        
        # 提取表格中的占位符
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        matches = re.findall(pattern, paragraph.text)
                        placeholders.update(matches)
        
        return list(placeholders)


# ------------------------------------------------------------------------------
# PDF模板处理器
# 功能: 处理PDF格式的模板文件
# 设计思路: 使用reportlab库生成PDF，支持中文字体
# ------------------------------------------------------------------------------
class PdfTemplateProcessor(BaseTemplateProcessor):
    """
    PDF模板处理器
    
    处理PDF格式的模板文件，使用reportlab库生成PDF输出。
    
    属性说明:
        template_content: 模板内容(文本形式)
        filled_content: 填充后的内容
    
    依赖:
        需要安装reportlab库: pip install reportlab
    
    特点:
        - 支持中文字体渲染
        - 自动分页
        - 如果reportlab不可用，自动降级为txt格式
    """
    
    def __init__(self):
        """初始化PDF模板处理器"""
        self.template_content = ""  # 模板内容
        self.filled_content = ""  # 填充后内容
    
    def load_template(self, template_path: str) -> bool:
        """
        加载PDF模板文件
        
        参数:
            template_path: 模板文件路径
        
        返回值:
            bool: 加载是否成功
        
        注意:
            当前实现将PDF作为文本读取，适合纯文本PDF模板。
        """
        try:
            with open(template_path, 'r', encoding='utf-8') as f:
                self.template_content = f.read()
            return True
        except Exception as e:
            return False
    
    def fill_template(self, data: Dict[str, Any]) -> str:
        """
        填充PDF模板
        
        参数:
            data: 键值对数据字典
        
        返回值:
            str: 填充后的文本内容
        """
        self.filled_content = self.template_content
        
        # 替换数据字段
        for key, value in data.items():
            placeholder = "{{" + key + "}}"
            self.filled_content = self.filled_content.replace(placeholder, str(value) if value else "")
        
        # 自动填充日期和时间
        self.filled_content = self.filled_content.replace("{{date}}", datetime.now().strftime("%Y-%m-%d"))
        self.filled_content = self.filled_content.replace("{{time}}", datetime.now().strftime("%H:%M:%S"))
        
        return self.filled_content
    
    def save(self, output_path: str) -> bool:
        """
        保存为PDF文件
        
        参数:
            output_path: 输出文件路径
        
        返回值:
            bool: 保存是否成功
        
        实现说明:
            1. 尝试使用reportlab生成PDF
            2. 如果reportlab不可用，降级保存为txt文件
            3. 支持中文字体渲染
        """
        try:
            # 导入reportlab相关模块
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.pdfgen import canvas
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # 创建PDF画布
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            # 尝试注册中文字体
            try:
                font_path = self._find_chinese_font()
                if font_path:
                    pdfmetrics.registerFont(TTFont('Chinese', font_path))
                    c.setFont('Chinese', 12)
                else:
                    c.setFont('Helvetica', 12)  # 使用默认字体
            except:
                c.setFont('Helvetica', 12)  # 字体注册失败时使用默认字体
            
            # 逐行绘制文本
            lines = self.filled_content.split('\n')
            y_position = height - 2 * cm  # 从页面顶部开始
            
            for line in lines:
                # 检查是否需要分页
                if y_position < 2 * cm:
                    c.showPage()  # 创建新页面
                    y_position = height - 2 * cm
                
                c.drawString(2 * cm, y_position, line)
                y_position -= 0.5 * cm  # 行间距0.5厘米
            
            c.save()
            return True
        except ImportError:
            # reportlab不可用，降级保存为txt
            return self._save_as_txt(output_path.replace('.pdf', '.txt'))
        except Exception as e:
            return False
    
    def _find_chinese_font(self) -> Optional[str]:
        """
        查找系统中的中文字体
        
        返回值:
            Optional[str]: 字体文件路径，未找到返回None
        
        实现说明:
            按优先级查找常见的中文字体文件:
            - Windows: 黑体、微软雅黑、宋体
            - Linux: Droid字体
            - macOS: 苹方字体
        """
        font_paths = [
            "C:/Windows/Fonts/simhei.ttf",  # Windows黑体
            "C:/Windows/Fonts/msyh.ttc",  # Windows微软雅黑
            "C:/Windows/Fonts/simsun.ttc",  # Windows宋体
            "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",  # Linux
            "/System/Library/Fonts/PingFang.ttc"  # macOS
        ]
        for path in font_paths:
            if os.path.exists(path):
                return path
        return None
    
    def _save_as_txt(self, output_path: str) -> bool:
        """
        降级保存为文本文件
        
        参数:
            output_path: 输出文件路径
        
        返回值:
            bool: 保存是否成功
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(self.filled_content)
            return True
        except:
            return False
    
    def get_placeholders(self) -> List[str]:
        """
        提取模板中的占位符
        
        返回值:
            List[str]: 占位符名称列表
        """
        pattern = r'\{\{(\w+)\}\}'
        return re.findall(pattern, self.template_content)


# ------------------------------------------------------------------------------
# 模板处理器管理类
# 功能: 统一管理不同格式的模板处理器，提供统一的处理接口
# 设计思路: 工厂模式，根据文件格式选择对应的处理器
# ------------------------------------------------------------------------------
class TemplateProcessor:
    """
    模板处理器管理类
    
    负责管理不同格式的模板处理器，提供统一的模板处理接口。
    
    属性说明:
        template_dir: 模板文件目录
        output_dir: 输出文件目录
        PROCESSORS: 格式到处理器类的映射字典
    
    主要方法:
        - process_template: 处理单个模板
        - get_template_placeholders: 获取模板占位符
        - batch_process: 批量处理模板
        - create_sample_template: 创建示例模板
    """
    
    # 格式到处理器类的映射
    PROCESSORS = {
        'txt': TxtTemplateProcessor,
        'docx': DocxTemplateProcessor,
        'pdf': PdfTemplateProcessor
    }
    
    def __init__(self, template_dir: str = "templates", output_dir: str = "output"):
        """
        初始化模板处理器
        
        参数:
            template_dir: 模板文件目录
            output_dir: 输出文件目录
        """
        self.template_dir = template_dir
        self.output_dir = output_dir
        self._ensure_directories()
    
    def _ensure_directories(self):
        """
        确保工作目录存在
        
        功能: 检查并创建模板目录和输出目录
        """
        os.makedirs(self.template_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)
    
    def _get_processor(self, file_format: str) -> Optional[BaseTemplateProcessor]:
        """
        获取指定格式的处理器实例
        
        参数:
            file_format: 文件格式(txt/docx/pdf)
        
        返回值:
            Optional[BaseTemplateProcessor]: 处理器实例，不支持的格式返回None
        """
        processor_class = self.PROCESSORS.get(file_format.lower())
        if processor_class:
            return processor_class()
        return None
    
    def process_template(
        self,
        template_path: str,
        data: Dict[str, Any],
        output_filename: str = None,
        output_format: str = None
    ) -> TemplateResult:
        """
        处理模板文件
        
        参数:
            template_path: 模板文件路径
            data: 填充数据字典
            output_filename: 输出文件名(可选)，未指定则自动生成
            output_format: 输出格式(可选)，未指定则与模板格式相同
        
        返回值:
            TemplateResult: 处理结果对象
        
        处理流程:
            1. 检查模板文件是否存在
            2. 确定输出格式
            3. 获取对应处理器
            4. 加载模板
            5. 填充数据
            6. 保存输出文件
        """
        # 检查模板文件
        if not os.path.exists(template_path):
            return TemplateResult(False, error_message=f"Template file not found: {template_path}")
        
        # 确定输出格式
        if output_format is None:
            output_format = os.path.splitext(template_path)[1][1:].lower()
        
        # 获取处理器
        processor = self._get_processor(output_format)
        if processor is None:
            return TemplateResult(False, error_message=f"Unsupported format: {output_format}")
        
        # 加载模板
        if not processor.load_template(template_path):
            return TemplateResult(False, error_message=f"Failed to load template: {template_path}")
        
        # 填充数据
        processor.fill_template(data)
        
        # 生成输出文件名
        if output_filename is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"output_{timestamp}.{output_format}"
        
        output_path = os.path.join(self.output_dir, output_filename)
        
        # 保存文件
        if processor.save(output_path):
            return TemplateResult(True, output_path=output_path)
        else:
            return TemplateResult(False, error_message=f"Failed to save output: {output_path}")
    
    def get_template_placeholders(self, template_path: str) -> List[str]:
        """
        获取模板中的占位符列表
        
        参数:
            template_path: 模板文件路径
        
        返回值:
            List[str]: 占位符名称列表
        
        实现说明:
            根据文件格式选择处理器，加载模板后提取占位符。
        """
        if not os.path.exists(template_path):
            return []
        
        file_format = os.path.splitext(template_path)[1][1:].lower()
        processor = self._get_processor(file_format)
        
        if processor and processor.load_template(template_path):
            return processor.get_placeholders()
        
        return []
    
    def batch_process(
        self,
        template_path: str,
        data_list: List[Dict[str, Any]],
        output_prefix: str = "output"
    ) -> List[TemplateResult]:
        """
        批量处理模板
        
        参数:
            template_path: 模板文件路径
            data_list: 数据字典列表，每个字典对应一个输出文件
            output_prefix: 输出文件名前缀
        
        返回值:
            List[TemplateResult]: 处理结果列表
        
        实现说明:
            遍历数据列表，为每组数据生成一个输出文件。
            文件名格式: 前缀_序号_时间戳.扩展名
        """
        results = []
        for i, data in enumerate(data_list):
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_filename = f"{output_prefix}_{i+1}_{timestamp}.{os.path.splitext(template_path)[1][1:]}"
            
            result = self.process_template(template_path, data, output_filename)
            results.append(result)
        
        return results
    
    def create_sample_template(self, template_name: str = "sample_template.txt") -> str:
        """
        创建示例模板文件
        
        参数:
            template_name: 模板文件名
        
        返回值:
            str: 创建的模板文件路径
        
        功能说明:
            创建一个包含常用占位符的示例模板，
            供用户参考和修改。
        """
        template_path = os.path.join(self.template_dir, template_name)
        
        # 示例模板内容
        sample_content = """====================================
              文档模板
====================================

标题: {{title}}
日期: {{date}}
时间: {{time}}

尊敬的 {{name}}:

您好！

根据您的请求，我们已处理以下内容：
{{content}}

联系方式: {{contact}}
地址: {{address}}

此致
敬礼！

====================================
"""
        
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(sample_content)
        
        return template_path


# ------------------------------------------------------------------------------
# 模板处理器工厂函数
# 功能: 创建模板处理器实例的便捷方法
# ------------------------------------------------------------------------------
def create_template_processor(template_dir: str = None, output_dir: str = None) -> TemplateProcessor:
    """
    创建模板处理器实例
    
    参数:
        template_dir: 模板目录(可选)，未指定则使用全局配置
        output_dir: 输出目录(可选)，未指定则使用全局配置
    
    返回值:
        TemplateProcessor: 配置好的模板处理器实例
    
    使用场景:
        快速创建模板处理器，自动使用全局配置中的目录设置。
    """
    if template_dir is None or output_dir is None:
        from config import CONFIG
        template_dir = template_dir or CONFIG.template.template_dir
        output_dir = output_dir or CONFIG.template.output_dir
    
    return TemplateProcessor(template_dir, output_dir)
