# ==============================================================================
# 文件名称: card_generator.py
# 功能描述: 席卡生成核心模块，提取公共的席卡生成逻辑供GUI和API共用
# 创建日期: 2026-07-02
# 作    者: 戒者有八
# 版本: 2.0.0
# ==============================================================================
"""
席卡生成核心模块

将 gui.py 和 server.py 中重复的席卡生成逻辑提取到此模块，
包括占位符替换、PDF转换、PDF合并、质量报告生成等功能。

支持模板:
    - 席卡模板v4.docx: {活动名称}、{姓名/公司} 等占位符
    - 席卡模板v5.docx: {text} 占位符（简化版，110pt大字号）
"""

import os
import re
import logging
from datetime import datetime
from typing import Dict, Any, List, Optional


# ------------------------------------------------------------------------------
# 文件名安全处理
# ------------------------------------------------------------------------------
def sanitize_filename(name: str) -> str:
    """处理文件名，移除非法字符"""
    illegal_chars = r'[<>:"/\\|?*]'
    safe_name = re.sub(illegal_chars, '_', name)
    return safe_name[:50] if len(safe_name) > 50 else safe_name


# ------------------------------------------------------------------------------
# 占位符替换（核心逻辑）
# ------------------------------------------------------------------------------
def replace_placeholder_in_paragraph(paragraph, display_text: str, event_name: str,
                                      display_type: str, info) -> None:
    """
    替换段落中的占位符并设置字体格式

    支持的占位符:
        v4模板: {活动名称}、{姓名/公司}、{{event}}、{{display}}、{{name}}、{{company}}、{{position}}、{{date}}
        v5模板: {text}（统一替换为display_text，保持模板原有字体格式）

    字体规则:
        - 活动名称: 楷体, 16pt
        - 姓名: 楷体, 72pt（两字姓名中间插入全角空格）
        - 公司名: 楷体, 36pt
        - {text}: 保持模板原有字体格式（v5模板已内置110pt楷体）
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn

    def format_name(text):
        if not text:
            return text
        if len(text) == 2 and all('\u4e00' <= c <= '\u9fff' for c in text):
            return text[0] + '\u3000' + text[1]
        return text

    is_company = (display_type == "company")
    name_font_size = 36 if is_company else 72

    formatted_display = display_text if is_company else format_name(display_text)
    formatted_name = info.name if is_company else format_name(info.name)

    # 标准占位符配置: (占位符, 替换值, 字体名称, 字号)
    placeholder_configs = [
        ('{活动名称}', event_name, '楷体', 16),
        ('{{event}}', event_name, '楷体', 16),
        ('{姓名/公司}', formatted_display, '楷体', name_font_size),
        ('{{display}}', formatted_display, '楷体', name_font_size),
        ('{{name}}', formatted_name, '楷体', name_font_size),
        ('{{company}}', info.company, '楷体', 36),
        ('{{position}}', info.position, None, None),
        ('{{date}}', datetime.now().strftime("%Y-%m-%d"), None, None),
    ]

    full_text = ''.join([run.text for run in paragraph.runs])

    # --- v5模板特殊处理: {text} 占位符保持模板原有字体格式 ---
    if '{text}' in full_text:
        _replace_text_placeholder_v5(paragraph, full_text, formatted_display)
        return

    # --- v4模板标准处理 ---
    found_placeholders = []
    for placeholder, value, font_name, font_size in placeholder_configs:
        if placeholder in full_text:
            found_placeholders.append((placeholder, value, font_name, font_size))

    if not found_placeholders:
        return

    for run in paragraph.runs:
        run.text = ""

    remaining_text = full_text

    for placeholder, value, font_name, font_size in found_placeholders:
        if placeholder in remaining_text:
            parts = remaining_text.split(placeholder, 1)

            if parts[0]:
                if paragraph.runs:
                    paragraph.runs[-1].text += parts[0]
                else:
                    paragraph.add_run(parts[0])

            if value:
                new_run = paragraph.add_run(str(value))
                if font_name:
                    try:
                        new_run.font.name = font_name
                        r = new_run._element
                        rPr = r.get_or_add_rPr()
                        rFonts = rPr.get_or_add_rFonts()
                        rFonts.set(qn('w:eastAsia'), font_name)
                    except Exception:
                        pass
                if font_size:
                    try:
                        new_run.font.size = Pt(font_size)
                    except Exception:
                        pass

            remaining_text = parts[1] if len(parts) > 1 else ""

    if remaining_text:
        if paragraph.runs:
            paragraph.runs[-1].text += remaining_text
        else:
            paragraph.add_run(remaining_text)


def _replace_text_placeholder_v5(paragraph, full_text: str, display_text: str) -> None:
    """
    处理v5模板的 {text} 占位符

    v5模板已内置字体格式（楷体 110pt），替换时保持原有run的格式不变，
    仅替换文本内容。
    """
    placeholder = '{text}'

    if placeholder not in full_text:
        return

    # 找到包含占位符的run
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, display_text)
            return

    # 如果占位符跨run分布，合并后替换
    if placeholder in full_text:
        # 清空所有run
        for run in paragraph.runs:
            run.text = ""
        # 在第一个run中写入替换后的文本，保持第一个run的格式
        if paragraph.runs:
            paragraph.runs[0].text = full_text.replace(placeholder, display_text)
        else:
            paragraph.add_run(full_text.replace(placeholder, display_text))


# ------------------------------------------------------------------------------
# Word转PDF
# ------------------------------------------------------------------------------
def convert_docx_to_pdf(docx_path: str, pdf_path: str) -> bool:
    """将Word文档转换为PDF，优先使用win32com，备选spire.doc"""
    try:
        import win32com.client
        word = win32com.client.Dispatch('Word.Application')
        word.Visible = False
        try:
            doc = word.Documents.Open(os.path.abspath(docx_path))
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
            doc.Close()
            return True
        finally:
            word.Quit()
    except Exception:
        try:
            from spire.doc import Document as SpireDocument
            doc = SpireDocument()
            doc.LoadFromFile(docx_path)
            doc.SaveToFile(pdf_path, 1)
            doc.Close()
            return True
        except Exception:
            return False


# ------------------------------------------------------------------------------
# PDF合并
# ------------------------------------------------------------------------------
def merge_pdfs(pdf_files: list, output_path: str) -> bool:
    """合并多个PDF文件为一个"""
    try:
        from PyPDF2 import PdfWriter, PdfReader
        merger = PdfWriter()
        for pdf_file in pdf_files:
            with open(pdf_file, 'rb') as f:
                reader = PdfReader(f)
                for page in reader.pages:
                    merger.add_page(page)
        with open(output_path, 'wb') as f:
            merger.write(f)
        return True
    except Exception:
        return False


# ------------------------------------------------------------------------------
# 文档质量验证
# ------------------------------------------------------------------------------
def validate_card_document(filepath: str, expected_content: str,
                           display_type: str = "name") -> dict:
    """验证席卡文档质量：检查是否为空白或内容不完整"""
    try:
        from docx import Document
        doc = Document(filepath)

        all_text = []
        for paragraph in doc.paragraphs:
            all_text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)

        full_text = ''.join(all_text)

        if not full_text.strip():
            return {'valid': False, 'reason': '文档内容为空'}

        if len(doc.paragraphs) == 0 and len(doc.tables) == 0:
            return {'valid': False, 'reason': '文档无段落或表格'}

        if expected_content:
            if expected_content in full_text:
                return {'valid': True, 'reason': '验证通过'}

            if display_type == "name" and len(expected_content) == 2:
                formatted = expected_content[0] + '\u3000' + expected_content[1]
                if formatted in full_text:
                    return {'valid': True, 'reason': '验证通过'}

            normalized = full_text.replace('\u3000', '').replace(' ', '')
            if expected_content in normalized:
                return {'valid': True, 'reason': '验证通过'}

            return {'valid': False, 'reason': f'文档未包含期望内容: {expected_content}'}

        return {'valid': True, 'reason': '验证通过'}

    except Exception as e:
        return {'valid': False, 'reason': f'验证过程出错: {str(e)}'}


# ------------------------------------------------------------------------------
# 质量报告生成
# ------------------------------------------------------------------------------
def generate_quality_report(report_path: str, log_data: dict,
                            generated_files: list, failed_persons: list) -> None:
    """生成席卡生成质量报告"""
    lines = [
        "=" * 60,
        "席卡生成质量报告".center(50),
        "=" * 60,
        "",
        f"生成时间: {log_data.get('start_time', 'N/A')}",
        f"活动名称: {log_data.get('event_name', '未指定')}",
        f"显示类型: {'公司名' if log_data.get('display_type') == 'company' else '姓名'}",
        f"总人数: {log_data.get('total_persons', 0)}",
        f"成功生成: {len([f for f in generated_files if f.get('valid', False)])}",
        f"验证失败: {len([f for f in generated_files if not f.get('valid', False)])}",
        f"PDF转换: {log_data.get('pdf_count', 0)}",
        "",
        f"输出目录: {log_data.get('output_dir', 'N/A')}",
        f"Word目录: {log_data.get('word_dir', 'N/A')}",
        f"PDF目录: {log_data.get('pdf_dir', 'N/A')}",
        "",
    ]

    if log_data.get('pdf_combined_path'):
        lines.extend([
            f"PDF合集: {os.path.basename(log_data['pdf_combined_path'])}",
            "",
        ])

    lines.extend(["-" * 60, "已生成文件列表:", "-" * 60])

    for f in generated_files:
        status = "✓ 有效" if f.get('valid', False) else f"✗ {f.get('reason', '无效')}"
        pdf_status = " (已转PDF)" if f.get('pdf_filepath') else ""
        lines.append(f"  {f['filename']}{pdf_status} - {status}")

    if failed_persons:
        lines.extend(["", "-" * 60, "失败记录:", "-" * 60])
        for name in failed_persons:
            lines.append(f"  • {name}")

    if log_data.get('warnings'):
        lines.extend(["", "-" * 60, "警告信息:", "-" * 60])
        for w in log_data['warnings']:
            lines.append(f"  ⚠ {w}")

    if log_data.get('errors'):
        lines.extend(["", "-" * 60, "错误信息:", "-" * 60])
        for e in log_data['errors']:
            lines.append(f"  ✗ {e}")

    lines.extend([
        "",
        "=" * 60,
        f"报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        "=" * 60,
    ])

    with open(report_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


# ------------------------------------------------------------------------------
# 席卡生成主流程
# ------------------------------------------------------------------------------
def generate_cards(infos, template_path: str, event_name: str,
                   display_type: str, output_base_dir: str) -> dict:
    """
    批量生成席卡（核心流程）

    参数:
        infos: PersonInfo对象列表
        template_path: 模板文件路径
        event_name: 活动名称
        display_type: 显示类型("name"或"company")
        output_base_dir: 输出基础目录

    返回:
        dict: 包含 success, count, output_dir, files, failed 等信息
    """
    log_data = {
        'start_time': datetime.now().isoformat(),
        'event_name': event_name,
        'display_type': display_type,
        'total_persons': len(infos),
        'errors': [],
        'warnings': [],
        'generated_files': []
    }

    try:
        from docx import Document

        if not os.path.exists(template_path):
            return {'success': False, 'error': f'模板文件不存在: {template_path}', 'log': log_data}

        # 创建输出目录结构
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        event_suffix = f"_{event_name}" if event_name else ""
        folder_name = f"席卡{event_suffix}_{timestamp}"
        output_dir = os.path.join(output_base_dir, folder_name)
        word_dir = os.path.join(output_dir, "word")
        pdf_dir = os.path.join(output_dir, "pdf")
        os.makedirs(word_dir, exist_ok=True)
        os.makedirs(pdf_dir, exist_ok=True)

        log_data['output_dir'] = output_dir
        log_data['word_dir'] = word_dir
        log_data['pdf_dir'] = pdf_dir

        generated_files = []
        failed_persons = []
        card_count = 0
        pdf_files = []

        for i, info in enumerate(infos):
            if display_type == "company":
                display_text = info.company if info.company else info.name
            else:
                display_text = info.name

            if not display_text:
                log_data['warnings'].append(f"第{i+1}条记录无有效显示内容，已跳过")
                continue

            try:
                doc = Document(template_path)

                for paragraph in doc.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, display_text, event_name, display_type, info)

                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_placeholder_in_paragraph(paragraph, display_text, event_name, display_type, info)

                date_str = datetime.now().strftime("%Y%m%d")
                safe_name = sanitize_filename(display_text)
                filename = f"席卡_{safe_name}_{date_str}.docx"
                filepath = os.path.join(word_dir, filename)

                counter = 1
                while os.path.exists(filepath):
                    filename = f"席卡_{safe_name}_{date_str}_{counter}.docx"
                    filepath = os.path.join(word_dir, filename)
                    counter += 1

                doc.save(filepath)

                validation = validate_card_document(filepath, display_text, display_type)

                if validation['valid']:
                    pdf_filename = filename.replace('.docx', '.pdf')
                    pdf_filepath = os.path.join(pdf_dir, pdf_filename)

                    if convert_docx_to_pdf(filepath, pdf_filepath):
                        pdf_files.append(pdf_filepath)
                        generated_files.append({
                            'filename': filename, 'filepath': filepath,
                            'pdf_filepath': pdf_filepath, 'name': display_text, 'valid': True
                        })
                    else:
                        log_data['warnings'].append(f"文件 {filename} 转换为PDF失败")
                        generated_files.append({
                            'filename': filename, 'filepath': filepath,
                            'name': display_text, 'valid': True, 'warning': 'PDF转换失败'
                        })
                    card_count += 1
                else:
                    log_data['warnings'].append(f"文件 {filename} 验证失败: {validation['reason']}")
                    generated_files.append({
                        'filename': filename, 'filepath': filepath,
                        'name': display_text, 'valid': False, 'reason': validation['reason']
                    })
                    failed_persons.append(display_text)

            except Exception as e:
                error_msg = f"生成 {display_text} 的席卡失败: {str(e)}"
                log_data['errors'].append(error_msg)
                failed_persons.append(display_text)

        # PDF合集
        pdf_combined_path = None
        if pdf_files:
            pdf_combined_path = os.path.join(output_dir, f"席卡合集_{timestamp}.pdf")
            if not merge_pdfs(pdf_files, pdf_combined_path):
                log_data['warnings'].append("PDF合集生成失败")
                pdf_combined_path = None

        if pdf_combined_path:
            log_data['pdf_combined_path'] = pdf_combined_path

        # 质量报告
        report_path = os.path.join(output_dir, "生成报告.txt")
        log_data['end_time'] = datetime.now().isoformat()
        log_data['generated_count'] = card_count
        log_data['failed_count'] = len(failed_persons)
        log_data['pdf_count'] = len(pdf_files)

        generate_quality_report(report_path, log_data, generated_files, failed_persons)

        return {
            'success': True,
            'count': card_count,
            'output_dir': output_dir,
            'word_dir': word_dir,
            'pdf_dir': pdf_dir,
            'files': generated_files,
            'failed': failed_persons,
            'log': log_data,
            'report_path': report_path,
            'pdf_combined_path': pdf_combined_path
        }

    except ImportError:
        return {'success': False, 'error': '请先安装python-docx库: pip install python-docx', 'log': log_data}
    except Exception as e:
        import traceback
        log_data['errors'].append(str(e))
        log_data['traceback'] = traceback.format_exc()
        return {'success': False, 'error': str(e), 'log': log_data}
